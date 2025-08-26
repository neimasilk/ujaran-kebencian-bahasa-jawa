#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script untuk mengekstrak dan membandingkan konten dari dua file Word (.docx)
untuk menyamakan draft paper dengan template jurnal.

Author: Assistant
Date: 2024
"""

import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def extract_text_content(doc_path):
    """
    Mengekstrak konten teks dari file Word (.docx)
    
    Args:
        doc_path (str): Path ke file .docx
        
    Returns:
        dict: Dictionary berisi informasi dokumen
    """
    try:
        doc = Document(doc_path)
        
        # Ekstrak paragraf
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    'text': para.text.strip(),
                    'style': para.style.name if para.style else 'Normal',
                    'alignment': para.alignment
                })
        
        # Ekstrak tabel
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        # Ekstrak styles yang tersedia
        styles = []
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                styles.append({
                    'name': style.name,
                    'type': 'paragraph'
                })
        
        return {
            'paragraphs': paragraphs,
            'tables': tables,
            'styles': styles,
            'total_paragraphs': len(paragraphs),
            'total_tables': len(tables)
        }
        
    except Exception as e:
        print(f"Error membaca file {doc_path}: {str(e)}")
        return None

def analyze_document_structure(content, doc_name):
    """
    Menganalisis struktur dokumen
    
    Args:
        content (dict): Konten dokumen hasil ekstraksi
        doc_name (str): Nama dokumen
    """
    print(f"\n{'='*60}")
    print(f"ANALISIS STRUKTUR: {doc_name}")
    print(f"{'='*60}")
    
    if not content:
        print("Gagal menganalisis dokumen.")
        return
    
    print(f"Total paragraf: {content['total_paragraphs']}")
    print(f"Total tabel: {content['total_tables']}")
    
    # Analisis styles yang digunakan
    print("\nStyles yang tersedia:")
    for style in content['styles'][:10]:  # Tampilkan 10 style pertama
        print(f"  - {style['name']}")
    
    # Analisis struktur berdasarkan style
    print("\nStruktur dokumen berdasarkan style:")
    style_count = {}
    for para in content['paragraphs']:
        style = para['style']
        style_count[style] = style_count.get(style, 0) + 1
    
    for style, count in sorted(style_count.items()):
        print(f"  {style}: {count} paragraf")
    
    # Tampilkan beberapa paragraf pertama
    print("\nBeberapa paragraf pertama:")
    for i, para in enumerate(content['paragraphs'][:5]):
        print(f"  {i+1}. [{para['style']}] {para['text'][:100]}...")
    
    # Analisis tabel jika ada
    if content['tables']:
        print(f"\nInformasi tabel:")
        for i, table in enumerate(content['tables']):
            print(f"  Tabel {i+1}: {len(table)} baris, {len(table[0]) if table else 0} kolom")
            if table:
                print(f"    Header: {table[0]}")

def compare_documents(draft_content, template_content):
    """
    Membandingkan struktur dan format antara draft dan template
    
    Args:
        draft_content (dict): Konten draft paper
        template_content (dict): Konten template jurnal
    """
    print(f"\n{'='*60}")
    print("PERBANDINGAN DOKUMEN")
    print(f"{'='*60}")
    
    if not draft_content or not template_content:
        print("Tidak dapat membandingkan dokumen karena ada yang gagal dibaca.")
        return
    
    # Bandingkan jumlah elemen
    print("\nPerbandingan jumlah elemen:")
    print(f"  Draft - Paragraf: {draft_content['total_paragraphs']}, Tabel: {draft_content['total_tables']}")
    print(f"  Template - Paragraf: {template_content['total_paragraphs']}, Tabel: {template_content['total_tables']}")
    
    # Bandingkan styles
    draft_styles = {style['name'] for style in draft_content['styles']}
    template_styles = {style['name'] for style in template_content['styles']}
    
    print("\nPerbandingan styles:")
    print(f"  Draft memiliki {len(draft_styles)} styles")
    print(f"  Template memiliki {len(template_styles)} styles")
    
    # Styles yang ada di template tapi tidak di draft
    missing_styles = template_styles - draft_styles
    if missing_styles:
        print(f"\n  Styles yang perlu ditambahkan ke draft:")
        for style in sorted(missing_styles):
            print(f"    - {style}")
    
    # Styles yang ada di draft tapi tidak di template
    extra_styles = draft_styles - template_styles
    if extra_styles:
        print(f"\n  Styles di draft yang tidak ada di template:")
        for style in sorted(extra_styles):
            print(f"    - {style}")

def generate_recommendations(draft_content, template_content):
    """
    Memberikan rekomendasi untuk menyesuaikan draft dengan template
    
    Args:
        draft_content (dict): Konten draft paper
        template_content (dict): Konten template jurnal
    """
    print(f"\n{'='*60}")
    print("REKOMENDASI PENYESUAIAN")
    print(f"{'='*60}")
    
    if not draft_content or not template_content:
        print("Tidak dapat memberikan rekomendasi.")
        return
    
    recommendations = []
    
    # Analisis struktur template untuk memberikan panduan
    template_structure = []
    for para in template_content['paragraphs'][:20]:  # Analisis 20 paragraf pertama
        if any(keyword in para['text'].lower() for keyword in ['abstract', 'abstrak', 'introduction', 'pendahuluan']):
            template_structure.append(f"Bagian: {para['text'][:50]}... [Style: {para['style']}]")
        elif any(keyword in para['text'].lower() for keyword in ['method', 'metodologi', 'hasil', 'result']):
            template_structure.append(f"Bagian: {para['text'][:50]}... [Style: {para['style']}]")
        elif any(keyword in para['text'].lower() for keyword in ['conclusion', 'kesimpulan', 'reference', 'daftar pustaka']):
            template_structure.append(f"Bagian: {para['text'][:50]}... [Style: {para['style']}]")
    
    print("\n1. STRUKTUR YANG DISARANKAN (berdasarkan template):")
    if template_structure:
        for i, section in enumerate(template_structure, 1):
            print(f"   {i}. {section}")
    else:
        print("   - Pastikan draft memiliki bagian: Abstract, Introduction, Method, Results, Conclusion, References")
    
    print("\n2. PENYESUAIAN FORMAT:")
    
    # Rekomendasi berdasarkan styles
    draft_styles = {style['name'] for style in draft_content['styles']}
    template_styles = {style['name'] for style in template_content['styles']}
    
    missing_styles = template_styles - draft_styles
    if missing_styles:
        print("   - Tambahkan styles berikut ke draft:")
        for style in sorted(missing_styles):
            print(f"     * {style}")
    
    # Rekomendasi berdasarkan tabel
    if template_content['total_tables'] > draft_content['total_tables']:
        print(f"   - Template memiliki {template_content['total_tables']} tabel, draft hanya {draft_content['total_tables']}")
        print("     Pertimbangkan menambahkan tabel sesuai kebutuhan")
    
    print("\n3. LANGKAH-LANGKAH PENYESUAIAN:")
    print("   1. Buka draft paper di Microsoft Word")
    print("   2. Buka template jurnal sebagai referensi")
    print("   3. Sesuaikan format heading dan subheading")
    print("   4. Pastikan margin, font, dan spacing sesuai template")
    print("   5. Sesuaikan format referensi dan sitasi")
    print("   6. Periksa format tabel dan gambar")
    print("   7. Pastikan struktur dokumen mengikuti template")

def main():
    """
    Fungsi utama untuk menjalankan analisis dan perbandingan
    """
    # Path ke file
    base_path = "d:/documents/ujaran-kebencian-bahasa-jawa/paper"
    draft_path = os.path.join(base_path, "draft paper.docx")
    template_path = os.path.join(base_path, "New_Template_Jurnal_JITK_Nusa_Mandiri_V11 No 1 August 2025.docx")
    
    print("SCRIPT PERBANDINGAN DOKUMEN WORD")
    print("=" * 60)
    print(f"Draft: {draft_path}")
    print(f"Template: {template_path}")
    
    # Periksa keberadaan file
    if not os.path.exists(draft_path):
        print(f"\nError: File draft tidak ditemukan: {draft_path}")
        return
    
    if not os.path.exists(template_path):
        print(f"\nError: File template tidak ditemukan: {template_path}")
        return
    
    print("\nMengekstrak konten dari kedua file...")
    
    # Ekstrak konten
    draft_content = extract_text_content(draft_path)
    template_content = extract_text_content(template_path)
    
    # Analisis masing-masing dokumen
    analyze_document_structure(draft_content, "DRAFT PAPER")
    analyze_document_structure(template_content, "TEMPLATE JURNAL")
    
    # Bandingkan dokumen
    compare_documents(draft_content, template_content)
    
    # Berikan rekomendasi
    generate_recommendations(draft_content, template_content)
    
    print(f"\n{'='*60}")
    print("ANALISIS SELESAI")
    print(f"{'='*60}")
    print("\nScript telah selesai menganalisis kedua dokumen.")
    print("Gunakan rekomendasi di atas untuk menyesuaikan draft dengan template.")

if __name__ == "__main__":
    main()