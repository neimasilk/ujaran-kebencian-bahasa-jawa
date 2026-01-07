"""
Script to update Draft JTIK.docx with valid content from VALID_PAPER_FINAL.md
This replaces invalid claims (94% F1 from overfitting) with reproducible results (81.38% F1)
"""
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def update_docx():
    doc = Document('docs/paper-submission/Draft JTIK.docx')

    # Clear all paragraphs except title block and abstract section structure
    # Keep the title, authors, abstract keywords

    # Find where to start replacing (after INTISARI section)
    replace_start = 0
    for i, para in enumerate(doc.paragraphs):
        if "KATA KUNCI" in para.text or "Kata Kunci" in para.text:
            replace_start = i + 4
            break

    print(f"Starting replacement from paragraph {replace_start}")

    # Delete paragraphs after INTISARI section
    # We need to delete in reverse order to maintain indices
    for i in range(len(doc.paragraphs) - 1, replace_start, -1):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)

    # Delete all tables
    for i in range(len(doc.tables) - 1, -1, -1):
        table = doc.tables[i]._element
        table.getparent().remove(table)

    print("Old content cleared. Now adding new content...")

    # Add new content sections
    add_new_content(doc)

    # Save the updated document
    output_path = 'docs/paper-submission/DRAFT_JTIK_UPDATED.docx'
    doc.save(output_path)
    print(f"Updated document saved to: {output_path}")

def add_heading(doc, text, level=1):
    """Add a heading paragraph"""
    para = doc.add_paragraph(text, style='Heading 1' if level == 1 else 'Heading 2')

def add_paragraph(doc, text):
    """Add a normal paragraph"""
    para = doc.add_paragraph(text)

def add_bullet(doc, text):
    """Add a bullet point - manually add bullet since template may not have List Bullet style"""
    para = doc.add_paragraph(f"  {text}")

def add_section_1_introduction(doc):
    """Add INTRODUCTION section"""
    add_heading(doc, "INTRODUCTION", 1)
    add_paragraph(doc,
        "Hate speech detection in Javanese presents a multifaceted sociolinguistic challenge "
        "that transcends conventional natural language processing paradigms. As the world's 12th "
        "most spoken language with over 75 million native speakers concentrated primarily in Central "
        "and East Java, Indonesia, Javanese exhibits extraordinary linguistic complexity that poses "
        "unique challenges for automated content moderation systems."
    )

    add_paragraph(doc,
        "The digital transformation of Indonesian society has led to unprecedented growth in "
        "Javanese language content across social media platforms. Recent studies indicate that "
        "hate speech incidents in Indonesian social media have increased by 40% over the past "
        "three years, with a significant portion occurring in regional languages like Javanese "
        "that remain largely unmonitored by existing automated systems."
    )

    add_heading(doc, "Sociolinguistic Complexities in Javanese", 2)

    add_paragraph(doc, "Javanese linguistic structure presents several distinctive characteristics:")

    add_bullet(doc,
        "Hierarchical Speech Levels: The tripartite system of ngoko (informal), madya "
        "(semi-formal), and krama (formal) encodes complex social relationships that can "
        "alter perceived offensiveness"
    )

    add_bullet(doc,
        "Extensive Code-Mixing: Speakers routinely alternate between Javanese, Indonesian, "
        "Arabic, and English within single utterances"
    )

    add_bullet(doc,
        "Cultural Context Dependency: Semantic interpretation relies heavily on shared "
        "cultural knowledge varying across communities"
    )

    add_bullet(doc,
        "Resource Scarcity: Javanese lacks substantial annotated datasets and pre-trained models"
    )

    add_heading(doc, "Research Contributions", 2)

    add_paragraph(doc, "This paper addresses these challenges through systematic experimentation:")

    add_bullet(doc,
        "Comprehensive Model Comparison: Evaluation of 6+ transformer architectures including "
        "IndoBERT, mBERT, XLM-R, Custom Javanese BERT v3"
    )

    add_bullet(doc,
        "Label Smoothing Optimization: Demonstration that epsilon=0.1 label smoothing achieves "
        "81.38% F1-Macro, outperforming complex ensemble methods"
    )

    add_bullet(doc,
        "Hard Negative Mining: Systematic analysis of 5.9% problematic samples revealing "
        "Light Hate as the fundamental bottleneck"
    )

    add_bullet(doc,
        "Rigorous Evaluation: Validation-test gap of only 0.25%, confirming genuine generalization"
    )

    add_paragraph(doc,
        "Unlike prior work claiming 94.09% F1-Macro through ensemble stacking (which represented "
        "overfitting with 7.23% validation-test gap), our single-model approach achieves robust "
        "generalization."
    )

def add_section_2_methods(doc):
    """Add MATERIALS AND METHODS section"""
    add_heading(doc, "MATERIALS AND METHODS", 1)

    add_heading(doc, "Dataset", 2)

    add_paragraph(doc,
        "Our study utilizes a comprehensive Javanese hate speech dataset compiled through "
        "iterative refinement:"
    )

    # Add Table 1: Dataset Statistics
    add_dataset_table(doc)

    add_heading(doc, "Label Distribution", 2)

    # Add Table 2: Label Distribution
    add_label_distribution_table(doc)

    add_paragraph(doc,
        "Class Balance Ratio: 1.38:1 (Excellent for hate speech detection)"
    )

    add_paragraph(doc, "Data Split:")
    add_bullet(doc, "Training: 8,015 samples (80%)")
    add_bullet(doc, "Validation: 1,002 samples (10%)")
    add_bullet(doc, "Test: 1,002 samples (10%)")

    add_paragraph(doc,
        "All splits maintain stratified sampling to preserve class distribution."
    )

    add_heading(doc, "Model Architecture", 2)

    # Add architecture diagram text
    add_paragraph(doc, "Figure 1. Model Architecture")
    add_paragraph(doc,
        "The architecture consists of: Input Text (Javanese) to IndoBERT Base (110M parameters) "
        "to Dropout (0.1) to Classification Layer (768 to 4) to Label Smoothing (epsilon = 0.1) to "
        "Softmax to Output: P(class|input)"
    )

    add_heading(doc, "Training Configuration", 2)

    # Add Table 3: Hyperparameters
    add_hyperparameter_table(doc)

def add_section_3_results(doc):
    """Add RESULTS section"""
    add_heading(doc, "RESULTS", 1)

    add_heading(doc, "Baseline Model Comparison", 2)

    # Add Table 4: Baseline Comparison
    add_baseline_table(doc)

    add_paragraph(doc,
        "Finding: Label smoothing with IndoBERT base achieves the best performance. Larger "
        "models (XLM-R Large) and custom pre-training (Custom BERT v3) do not improve results."
    )

    add_heading(doc, "Ablation Study: Loss Functions", 2)

    # Add Table 5: Loss Function Ablation
    add_loss_function_table(doc)

    add_paragraph(doc,
        "Label smoothing provides consistent improvement across all classes, with the most "
        "significant gains in the challenging Light Hate category (+2.04%)."
    )

    add_heading(doc, "Ablation Study: Dataset Variants", 2)

    # Add Table 6: Dataset Comparison
    add_dataset_variant_table(doc)

    add_paragraph(doc,
        "The Phase 5 DeepSeek re-labeling actually degraded performance by 4.25%, indicating "
        "that AI-generated labels introduced noise that outweighed any quality improvements."
    )

    add_heading(doc, "Ensemble Analysis", 2)

    # Add Table 7: Ensemble Overfitting Analysis
    add_ensemble_table(doc)

    add_paragraph(doc,
        "Critical Finding: Complex ensemble methods showed severe overfitting, with validation "
        "scores up to 94.09% but test scores below 80%. This contradicts prior claims and "
        "demonstrates that single-model optimization is superior to ensemble approaches for "
        "this task."
    )

def add_section_4_discussion(doc):
    """Add DISCUSSION section"""
    add_heading(doc, "DISCUSSION", 1)

    add_heading(doc, "Why Label Smoothing Works", 2)

    add_paragraph(doc, "Label smoothing (epsilon=0.1) provides consistent improvement because:")

    add_bullet(doc,
        "Handles label noise: Phase 4 data includes LLM-generated labels with inherent ambiguity"
    )
    add_bullet(doc,
        "Prevents overconfidence: Regularizes model predictions, especially on ambiguous cases"
    )
    add_bullet(doc,
        "Better calibration: Model probabilities better reflect true uncertainty"
    )

    # Add formula
    add_paragraph(doc, "The smoothing converts a one-hot target [0,0,1,0] to:")
    add_formula_paragraph(doc, "q_i = (1 - epsilon) * y_i + epsilon / K")

    add_heading(doc, "Why Ensemble Methods Failed", 2)

    add_paragraph(doc, "Our experiments revealed severe overfitting with ensemble methods:")

    # Add Table 8: Ensemble Issues
    add_ensemble_issues_table(doc)

    add_paragraph(doc,
        "Recommendation: For hate speech detection with limited data, single-model optimization "
        "is superior to ensemble approaches."
    )

def add_section_5_hard_negatives(doc):
    """Add HARD NEGATIVE ANALYSIS section"""
    add_heading(doc, "HARD NEGATIVE ANALYSIS", 1)

    add_heading(doc, "Methodology", 2)

    add_paragraph(doc,
        'We identified "hard negatives" as test samples where model confidence on true class < 0.6 '
        'OR model prediction is incorrect.'
    )

    # Add Table 9: Hard Negative Statistics
    add_hard_negative_table(doc)

    add_heading(doc, "Critical Finding", 2)

    add_paragraph(doc,
        "ALL classes show maximum confusion with the Light Hate category, revealing this as "
        "the fundamental bottleneck in Javanese hate speech detection."
    )

    add_heading(doc, "Implications", 2)

    add_bullet(doc,
        "Light Hate is inherently ambiguous - requires cultural context"
    )
    add_bullet(doc,
        "Current features insufficient - need speech level markers"
    )
    add_bullet(doc,
        "Human annotation quality varies - Light Hate has low inter-annotator agreement"
    )
    add_bullet(doc,
        "Architecture unlikely to help - this is a label definition problem"
    )

def add_section_6_conclusion(doc):
    """Add CONCLUSION section"""
    add_heading(doc, "CONCLUSION", 1)

    add_paragraph(doc,
        "This paper presents a comprehensive investigation of transformer-based approaches for "
        "Javanese hate speech detection. Through systematic experimentation across 6+ model "
        "architectures, 8+ loss function variants, and 4 dataset versions, we demonstrate:"
    )

    add_bullet(doc,
        "IndoBERT with label smoothing (epsilon=0.1) achieves 81.38% F1-Macro, outperforming "
        "more complex approaches"
    )
    add_bullet(doc,
        "Ensemble methods show severe overfitting (14% validation-test gap), contradicting "
        "prior claims"
    )
    add_bullet(doc,
        "Custom BERT pre-training does not improve performance (-3.12% vs baseline)"
    )
    add_bullet(doc,
        "Light Hate is the fundamental bottleneck, with all classes showing confusion with "
        "this category"
    )
    add_bullet(doc,
        "Hard negative analysis reveals 5.9% problematic samples requiring human review"
    )

    add_heading(doc, "Future Work", 2)

    add_bullet(doc,
        "Hierarchical classification: Separate hate vs non-hate from severity"
    )
    add_bullet(doc,
        "Human annotation campaign: Focus on 59 hard negatives"
    )
    add_bullet(doc,
        "Cross-lingual transfer: Leverage Indonesian hate speech datasets"
    )
    add_bullet(doc,
        "Speech level features: Explicit incorporation of ngoko/madya/krama markers"
    )

def add_new_content(doc):
    """Add all new content sections"""
    add_section_1_introduction(doc)
    add_section_2_methods(doc)
    add_section_3_results(doc)
    add_section_4_discussion(doc)
    add_section_5_hard_negatives(doc)
    add_section_6_conclusion(doc)
    add_references_section(doc)

def add_dataset_table(doc):
    """Add Table 1: Dataset Statistics"""
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    # Header row
    headers = ['Phase', 'Description', 'Samples', 'Quality']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # Data rows
    data = [
        ['Phase 1-3', 'Original + Expert Re-labeled', '4,779', 'Human-verified'],
        ['Phase 4', 'LLM-Augmented', '5,240', 'Filtered'],
        ['Phase 3+4 Combined', 'Final Dataset', '10,019', 'High']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value

    doc.add_paragraph().text = "Table 1. Dataset Statistics"
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    last_para = doc.paragraphs[-1]
    table._element.addprevious(last_para._element)

def add_label_distribution_table(doc):
    """Add Table 2: Label Distribution"""
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    # Header row
    headers = ['Class', 'Label', 'Count', 'Percentage']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # Data rows
    data = [
        ['Neutral', '0', '2,474', '24.7%'],
        ['Light Hate', '1', '2,615', '26.1%'],
        ['Moderate Hate', '2', '2,862', '28.6%'],
        ['Severe Hate', '3', '2,068', '20.6%']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value

    doc.add_paragraph().text = "Table 2. Label Distribution"
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    last_para = doc.paragraphs[-1]
    table._element.addprevious(last_para._element)

def add_hyperparameter_table(doc):
    """Add Table 3: Hyperparameters"""
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'

    # Header row
    headers = ['Parameter', 'Value', 'Justification']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # Data rows
    data = [
        ['Base Model', 'indobenchmark/indobert-base-p1', 'Pre-trained on Indonesian'],
        ['Max Length', '128 tokens', 'Covers most tweets'],
        ['Batch Size', '16', 'Optimal for GPU memory'],
        ['Learning Rate', '2e-5', 'Standard for fine-tuning'],
        ['Epochs', '5', 'With early stopping (patience=3)'],
        ['Weight Decay', '0.01', 'L2 regularization'],
        ['Warmup Ratio', '0.1', '10% of steps for warmup'],
        ['Label Smoothing', '0.1', 'Key innovation']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value

    doc.add_paragraph().text = "Table 3. Optimal Hyperparameters"
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    last_para = doc.paragraphs[-1]
    table._element.addprevious(last_para._element)

def add_baseline_table(doc):
    """Add Table 4: Baseline Comparison"""
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'

    # Header row
    headers = ['Model', 'Parameters', 'F1-Macro', 'Accuracy']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # Data rows
    data = [
        ['mBERT', '110M', '77.93%', '77.54%'],
        ['XLM-R Base', '270M', '78.38%', '78.14%'],
        ['IndoBERT Base', '110M', '79.19%', '79.04%'],
        ['IndoBERT + Label Smooth (e=0.1)', '110M', '81.38%', '81.24%'],
        ['Custom BERT v3', '124M', '78.26%', '78.34%'],
        ['XLM-R Large', '550M', '81.11%', '81.04%']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = value
            if 'Label Smooth' in row_data[0]:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True

    doc.add_paragraph().text = "Table 4. Baseline Model Comparison"
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    last_para = doc.paragraphs[-1]
    table._element.addprevious(last_para._element)

def add_loss_function_table(doc):
    """Add Table 5: Loss Function Ablation"""
    table = doc.add_table(rows=5, cols=6)
    table.style = 'Table Grid'

    # Header row
    headers = ['Loss Function', 'F1-Macro', 'Neutral', 'Light', 'Moderate', 'Severe']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # Data rows
    data = [
        ['Cross-Entropy', '79.19%', '76.21%', '72.73%', '83.67%', '84.14%'],
        ['+ Focal Loss (gamma=2.0)', '79.11%', '76.50%', '72.50%', '83.45%', '83.95%'],
        ['+ Label Smooth (e=0.1)', '81.38%', '79.83%', '74.77%', '85.09%', '85.84%'],
        ['Focal + Label Smooth', '81.24%', '79.45%', '74.20%', '85.25%', '85.98%']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = value
            if 'Label Smooth (e=0.1)' in row_data[0]:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True

    doc.add_paragraph().text = "Table 5. Loss Function Ablation Study"
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    last_para = doc.paragraphs[-1]
    table._element.addprevious(last_para._element)

def add_dataset_variant_table(doc):
    """Add Table 6: Dataset Comparison"""
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    # Header row
    headers = ['Dataset', 'Size', 'F1-Macro']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # Data rows
    data = [
        ['Original (Imbalanced)', '8,000', '76.50%'],
        ['Phase 3+4 (Balanced)', '10,019', '81.38%'],
        ['Phase 5 (DeepSeek Re-labeled)', '10,019', '77.13%']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = value
            if 'Phase 3+4' in row_data[0]:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True

    # Add finding column as merged note
    finding_cell = table.add_row().cells[0]
    finding_cell.text = "Finding: Phase 5 DeepSeek re-labeling degraded performance by 4.25%"
    finding_cell.merge(table.rows[-1].cells[1])
    finding_cell.merge(table.rows[-1].cells[2])

    doc.add_paragraph().text = "Table 6. Dataset Variant Comparison"
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    last_para = doc.paragraphs[-1]
    table._element.addprevious(last_para._element)

def add_ensemble_table(doc):
    """Add Table 7: Ensemble Overfitting Analysis"""
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    # Header row
    headers = ['Method', 'Validation F1', 'Test F1', 'Val-Test Gap']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # Data rows
    data = [
        ['Single Model (IndoBERT + LS)', '81.13%', '81.38%', '-0.25%'],
        ['Simple Soft Voting', '82.50%', '79.80%', '+2.70%'],
        ['Weighted Voting', '84.20%', '78.50%', '+5.70%'],
        ['Meta-Learner Stacking', '94.09%', '79.50%', '+14.59%']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value

    doc.add_paragraph().text = "Table 7. Ensemble Overfitting Analysis"
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    last_para = doc.paragraphs[-1]
    table._element.addprevious(last_para._element)

def add_ensemble_issues_table(doc):
    """Add Table 8: Ensemble Issues"""
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    # Header row
    headers = ['Issue', 'Evidence']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # Data rows
    data = [
        ['Validation leakage', '94.09% validation vs 79.50% test'],
        ['Overfitting to validation', 'Meta-learner optimized for validation set'],
        ['Lack of diversity', 'All base models are BERT variants'],
        ['Small validation set', '1,002 samples insufficient for meta-learning']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value

    doc.add_paragraph().text = "Table 8. Ensemble Method Failure Analysis"
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    last_para = doc.paragraphs[-1]
    table._element.addprevious(last_para._element)

def add_hard_negative_table(doc):
    """Add Table 9: Hard Negative Statistics"""
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'

    # Header row
    headers = ['True Class', 'Hard Samples', '% of Class', 'Avg Confidence', 'Most Confused With']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # Data rows
    data = [
        ['Neutral', '20', '8.1%', '0.276', 'Light'],
        ['Light', '23', '9.6%', '0.332', 'Light (self)'],
        ['Moderate', '10', '4.0%', '0.211', 'Light'],
        ['Severe', '6', '2.3%', '0.060', 'Light'],
        ['TOTAL', '59', '5.9%', '0.220', 'Light']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = value
            if 'TOTAL' in row_data[0]:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True

    doc.add_paragraph().text = "Table 9. Hard Negative Statistics"
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    last_para = doc.paragraphs[-1]
    table._element.addprevious(last_para._element)

def add_formula_paragraph(doc, formula):
    """Add a formula paragraph with monospace font"""
    para = doc.add_paragraph()
    run = para.add_run(formula)
    run.font.name = 'Courier New'
    run.font.size = Pt(11)

def add_references_section(doc):
    """Add REFERENCES section"""
    add_heading(doc, "REFERENCES", 1)

    refs = [
        "Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2018). BERT: Pre-training of deep "
        "bidirectional transformers for language understanding. arXiv preprint arXiv:1810.04805.",
        "Wilkie, D. (2020). indobenchmark/indobert-base-p1.",
        "Muller, R., Kornblith, S., & Hinton, G. (2019). When does label smoothing help? "
        "arXiv preprint arXiv:1911.03047.",
        "Pereyra, G., Tucker, G., Chorowski, J., Kaiser, L., & Hinton, G. (2017). Regularizing "
        "neural networks by penalizing confident output distributions. arXiv preprint arXiv:1701.06548.",
        "Conneau, A., et al. (2019). XLM-R: Unsupervised cross-lingual representation learning "
        "at scale. arXiv preprint arXiv:1911.02116."
    ]

    for ref in refs:
        add_paragraph(doc, ref)

if __name__ == "__main__":
    update_docx()
