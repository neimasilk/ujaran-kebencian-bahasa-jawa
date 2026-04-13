#!/usr/bin/env python3
"""
Generate Kinetik (SINTA 2) paper DOCX from template and experimental results.
Version 2: With methodological fixes (baselines, augmentation impact, examples, references).

Usage:
    python paper/generate_kinetik_paper.py

Input:
    - Template Kinetik Mendeley.docx (template with styles)
    - results/*.json (experimental results)
    - paper/figures/*.png (figures)

Output:
    - paper/paper_kinetik.docx (final paper ready for submission)
"""

import json
import statistics
from pathlib import Path
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# PATHS
# ============================================================
BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE = BASE_DIR / "Template Kinetik Mendeley.docx"
OUTPUT = BASE_DIR / "paper" / "paper_kinetik.docx"
RESULTS = BASE_DIR / "results"
FIGURES = BASE_DIR / "paper" / "figures"


def generate_manual_confusion_matrix(manual_only):
    """Generate confusion matrix figure from manual-only test data."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import seaborn as sns
    except ImportError:
        print("[WARN] matplotlib/seaborn not installed, skipping CM generation")
        return

    if not manual_only:
        return

    cm_data = manual_only["models"]["xlmr_large"]["manual_only"]["confusion_matrix"]
    cm = np.array(cm_data)
    labels = ["Not Hate", "Light", "Moderate", "Severe"]

    fig, ax = plt.subplots(figsize=(6, 5))
    sns.heatmap(
        cm, annot=True, fmt="d", cmap="Blues",
        xticklabels=labels, yticklabels=labels,
        ax=ax, cbar_kws={"shrink": 0.8}
    )
    ax.set_xlabel("Predicted", fontsize=11)
    ax.set_ylabel("Actual", fontsize=11)
    ax.set_title("XLM-RoBERTa Large — Manual-Only Test (451 samples)", fontsize=12)
    plt.tight_layout()

    out_path = FIGURES / "figure3_confusion_matrix_manual.png"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(str(out_path), dpi=150)
    plt.close(fig)
    print(f"  Generated confusion matrix: {out_path}")


# ============================================================
# DATA LOADING
# ============================================================
def load_json(name):
    path = RESULTS / f"{name}.json"
    if not path.exists():
        print(f"[WARN] Not found: {path}")
        return None
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def load_cleaning():
    with open(
        BASE_DIR / "data" / "cleaned" / "cleaning_report.json", encoding="utf-8"
    ) as f:
        return json.load(f)


# ============================================================
# HELPER FUNCTIONS
# ============================================================
def clear_body(doc):
    """Remove all paragraphs and tables from template body, keep sectPr."""
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def update_footers(doc):
    """Update footer citation from template placeholder to our paper's citation."""
    from xml.sax.saxutils import escape as xml_escape
    cite_text = (
        "Cite: Amien, M., Sijabat, D. R., & Kanthi, Y. A. (2026). "
        "Evaluating LLM-Augmented Transformer Models for Severity-Based "
        "Hate Speech Detection in Javanese. "
        "Kinetik: Game Technology, Information System, Computer Network, "
        "Computing, Electronics, and Control."
    )
    cite_xml = xml_escape(cite_text)

    def _make_footer_p():
        return parse_xml(
            f'<w:p {nsdecls("w")}>'
            f'<w:pPr>'
            f'<w:pStyle w:val="Footer"/>'
            f'<w:pBdr><w:top w:val="single" w:sz="4" w:space="1" w:color="auto"/></w:pBdr>'
            f'<w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
            f'<w:i/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
            f'</w:pPr>'
            f'<w:r><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
            f'<w:i/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
            f'<w:t xml:space="preserve">{cite_xml}</w:t></w:r>'
            f'</w:p>'
        )

    for section in doc.sections:
        # Main footer (odd/subsequent pages)
        # Clear all content (including hyperlinks with old DOI) and rebuild
        footer = section.footer
        if not footer.is_linked_to_previous:
            ft_el = footer._element
            for child in list(ft_el):
                ft_el.remove(child)
            ft_el.append(_make_footer_p())

        # First page footer — clear all content (drawings, shapes, text)
        # and replace with our citation
        fp_footer = section.first_page_footer
        if not fp_footer.is_linked_to_previous:
            fp_el = fp_footer._element
            for child in list(fp_el):
                fp_el.remove(child)
            fp_el.append(_make_footer_p())


def _try_style(doc, style_name, fallback="Normal"):
    """Return style_name if it exists in the document, else fallback."""
    try:
        doc.styles[style_name]
        return style_name
    except KeyError:
        return fallback


def add_body(doc, text="", bold=False, italic=False):
    """Add a body paragraph."""
    style = _try_style(doc, "Body kinetik")
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    return p


def add_heading_main(doc, text):
    """Add a main section heading (1., 2., 3., 4.)."""
    style = _try_style(doc, "Sub title 1 kinetik")
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_subheading(doc, text):
    """Add a sub-section heading (bold body text)."""
    style = _try_style(doc, "Body kinetik")
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = True
    p.paragraph_format.space_before = Pt(6)
    return p


def add_caption(doc, text):
    """Add a figure/table caption (centered, italic)."""
    style = _try_style(doc, "Figure table kinetik")
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    return p


def add_figure(doc, filename, caption_text, width=Inches(5.0)):
    """Add a figure with centered caption."""
    path = FIGURES / filename
    if not path.exists():
        add_body(doc, f"[Figure not found: {filename}]", italic=True)
        return
    doc.add_picture(str(path), width=width)
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption_text)


def fmt_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=9):
    """Format a table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table):
    """Set academic-style table borders (top, bottom, header separator)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        "</w:tblBorders>"
    )
    tblPr.append(borders)


def set_header_shading(table, n_cols, color="D9E2F3"):
    """Shade the header row."""
    for i in range(n_cols):
        cell = table.cell(0, i)
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
        )
        cell._element.get_or_add_tcPr().append(shading)


def create_table(doc, headers, rows, caption_text, col_widths=None):
    """Create a formatted table with caption above."""
    add_caption(doc, caption_text)

    n_cols = len(headers)
    n_rows = len(rows) + 1  # +1 for header
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header row
    for i, h in enumerate(headers):
        fmt_cell(table.cell(0, i), h, bold=True)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            bold = isinstance(val, str) and val.startswith("**")
            if bold:
                val = val.strip("*")
            fmt_cell(table.cell(r_idx + 1, c_idx), str(val), bold=bold)

    set_table_borders(table)
    set_header_shading(table, n_cols)

    # Spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)

    return table


# ============================================================
# PAPER SECTIONS
# ============================================================


def write_title(doc):
    style = _try_style(doc, "Title Kinetik")
    p = doc.add_paragraph(style=style)
    p.add_run(
        "Evaluating LLM-Augmented Transformer Models for "
        "Severity-Based Hate Speech Detection in Javanese"
    )


def write_authors(doc):
    style_auth = _try_style(doc, "Authors Kinetik")
    p = doc.add_paragraph(style=style_auth)
    p.add_run("Mukhlis Amien")
    r = p.add_run("*1")
    r.font.superscript = True
    p.add_run(", Daniel Rudiaman Sijabat")
    r = p.add_run("2")
    r.font.superscript = True
    p.add_run(", Yekti Asmoro Kanthi")
    r = p.add_run("3")
    r.font.superscript = True

    style_aff = _try_style(doc, "Afiliation Kinetik")
    p = doc.add_paragraph(style=style_aff)
    r = p.add_run("1,2")
    r.font.superscript = True
    p.add_run("Department of Informatics, ")
    r = p.add_run("3")
    r.font.superscript = True
    p.add_run("Department of Information System")

    p = doc.add_paragraph(style=style_aff)
    p.add_run("Universitas Bhinneka Nusantara, Malang, Indonesia")
    p = doc.add_paragraph(style=style_aff)
    p.add_run("*mukhlisramien@gmail.com")


def write_abstract(doc, comp, baselines, augmentation, manual_only=None, multiseed=None):
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    m_xl = comp["models"]["xlmr_large"]
    svm = baselines["models"]["svm_tfidf"] if baselines else None

    # Get manual-only F1 for primary reporting
    manual_f1 = None
    n_manual_test = 0
    if manual_only:
        manual_f1 = manual_only["models"]["xlmr_large"]["manual_only"]["f1_macro"]
        n_manual_test = manual_only["metadata"]["manual_test_samples"]

    # Compute multi-seed stats from data
    ms_text = ""
    if multiseed:
        stable = [r["test"]["f1_macro"] for r in multiseed["runs"] if r["seed"] != 1024]
        if len(stable) >= 2:
            ms_mean = statistics.mean(stable)
            ms_std = statistics.pstdev(stable)
            ms_text = f"{ms_mean:.2f}% \u00b1 {ms_std:.2f}%"

    # === 2-Column Table: Article Info (left) | Abstract (right) ===
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False

    # Set column widths (~30% left, ~70% right)
    for cell in table.columns[0].cells:
        cell.width = Cm(5.0)
    for cell in table.columns[1].cells:
        cell.width = Cm(12.0)

    # Remove table borders (Kinetik style — clean look)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>'
    )
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # --- LEFT CELL: Article Info ---
    left_cell = table.cell(0, 0)
    # Clear default paragraph
    left_cell.paragraphs[0].clear()

    style_ai = _try_style(doc, "Article info kinetik")

    # Article Info heading
    p = left_cell.paragraphs[0]
    p.style = _try_style(doc, "Article info heading kinetik")
    r = p.add_run("Article Info")
    r.bold = True

    # Article history
    p = left_cell.add_paragraph(style=style_ai)
    r = p.add_run("Article history:")
    r.bold = True

    for label in ["Received", "Revised", "Accepted", "Published"]:
        p = left_cell.add_paragraph(style=style_ai)
        p.add_run(f"{label}: [date]")

    # Keywords (alphabetical, no title-word overlap)
    p = left_cell.add_paragraph(style=style_ai)
    r = p.add_run("Keywords:")
    r.bold = True

    p = left_cell.add_paragraph(style=style_ai)
    p.add_run(
        "Data Augmentation\n"
        "Low-Resource NLP\n"
        "Pre-trained Language Model\n"
        "Severity Classification\n"
        "Text Classification"
    )

    # Corresponding Author
    p = left_cell.add_paragraph(style=style_ai)
    r = p.add_run("Corresponding Author:")
    r.bold = True

    p = left_cell.add_paragraph(style=style_ai)
    p.add_run(
        "Mukhlis Amien,\n"
        "Department of Informatics,\n"
        "Universitas Bhinneka Nusantara,\n"
        "Malang, Indonesia\n"
        "mukhlisramien@gmail.com"
    )

    # --- RIGHT CELL: Abstract ---
    right_cell = table.cell(0, 1)
    right_cell.paragraphs[0].clear()

    # Abstract heading
    p = right_cell.paragraphs[0]
    p.style = _try_style(doc, "Abstract heading kinetik")
    p.add_run("Abstract")

    # Abstract text (past tense per Kinetik guidelines)
    style_abs = _try_style(doc, "Abstract Kinetik")
    p = right_cell.add_paragraph(style=style_abs)
    p.add_run(
        "This study empirically evaluated transformer-based models for "
        "severity-based hate speech detection in Javanese, a low-resource "
        "language spoken by approximately 80 million people. A dataset of "
        "9,775 annotated samples was constructed through manual collection "
        "(46.4%) and LLM-based augmentation using DeepSeek-Coder-V2 (53.6%), "
        "with four severity levels: Not Hate Speech, Light, Moderate, and "
        "Severe. A comparative study was conducted across traditional "
        "baselines (SVM, Logistic Regression with TF-IDF) and transformer "
        "architectures (IndoBERT, XLM-RoBERTa Large). "
    )
    if manual_f1:
        p.add_run(
            f"On manually collected test data ({n_manual_test} samples), "
            f"XLM-RoBERTa Large achieved the best F1-Macro of {manual_f1:.2f}%"
        )
        if manual_only and "svm_tfidf" in manual_only["models"]:
            svm_m = manual_only["models"]["svm_tfidf"]["manual_only"]["f1_macro"]
            ib_m = manual_only["models"]["indobert"]["manual_only"]["f1_macro"]
            p.add_run(
                f", outperforming SVM + TF-IDF ({svm_m:.2f}%) and "
                f"IndoBERT ({ib_m:.2f}%)"
            )
        if augmentation:
            synth_f1 = augmentation["results"].get("full_model_on_synthetic_test", {}).get("f1_macro", 0)
            if synth_f1 > 0:
                p.add_run(
                    f". A key finding was the substantial performance gap between "
                    f"synthetic ({synth_f1:.2f}%) and manual ({manual_f1:.2f}%) "
                    "test data, revealing that LLM augmentation, while necessary "
                    "to prevent training collapse, introduced systematic "
                    "distributional differences that inflated standard "
                    "evaluation metrics. "
                )
            else:
                p.add_run(". ")
        else:
            p.add_run(". ")
    else:
        p.add_run(
            f"Experimental results demonstrated that XLM-RoBERTa Large achieved "
            f"the best performance with an F1-Macro of {m_xl['test']['f1_macro']}% "
            f"on the test set. "
        )
    if ms_text:
        p.add_run(
            "Multi-seed evaluation across five random seeds confirmed "
            f"model stability with a mean F1 of {ms_text} on the full test set "
            "across four stable seeds. "
        )
    p.add_run(
        "These findings provided empirical evidence on the effectiveness "
        "and limitations of LLM-augmented training for hate speech detection "
        "in low-resource regional languages."
    )


# ------------------------------------------------------------------
# SECTION 1: INTRODUCTION
# ------------------------------------------------------------------
def write_introduction(doc):
    add_heading_main(doc, "1. Introduction")

    # Para 1: Global context
    p = add_body(doc)
    p.add_run(
        "Hate speech on social media has become a pressing global concern, "
        "exacerbated by the rapid growth of online platforms [16,17]. Automated "
        "detection systems are essential for identifying and mitigating harmful "
        "content at scale. Comprehensive reviews have catalogued the evolution "
        "of detection approaches from traditional machine learning to "
        "transformer-based models [9,18], while recent work has also explored "
        "explainability in hate speech classification [19]. Despite this "
        "progress, most research has focused on high-resource languages such "
        "as English, leaving regional and low-resource languages significantly "
        "underserved."
    )

    # Para 2: Indonesia & Javanese
    p = add_body(doc)
    p.add_run(
        "Indonesia, with over 200 million internet users, faces significant "
        "challenges in combating hate speech across its diverse linguistic "
        "landscape. Hate speech detection for standard Indonesian has received "
        "considerable attention. Ibrohim and Budi [1] introduced a multi-label "
        "dataset of 13,069 Indonesian tweets and achieved an F1-score of "
        "71.31% using Bidirectional LSTM with FastText embeddings. Alfina et "
        "al. [2] developed one of the earliest Indonesian hate speech datasets "
        "with 713 samples and compared Naive Bayes with Support Vector "
        "Machine classifiers. However, regional languages remain largely "
        "understudied. Javanese, spoken by approximately 80 million people as "
        "the largest regional language in Indonesia, is classified as a "
        "low-resource language in the NLP context [10]."
    )

    # Para 3: Regional language NLP
    p = add_body(doc)
    p.add_run(
        "Research on NLP for Indonesian regional languages has gained momentum "
        "in recent years. Putri et al. [3] conducted a preliminary study on "
        "hate speech detection for Javanese and Sundanese tweets, though "
        "limited to binary classification with a small dataset. Cahyawijaya "
        "et al. [5] unified 137 datasets for Indonesian and 18+ regional "
        "languages through the NusaCrowd initiative, establishing benchmarks "
        "for zero-shot evaluation. IndoNLG [21] expanded resources for "
        "Indonesian natural language generation, while NusaX [22] provided a "
        "multilingual parallel sentiment dataset for 10 Indonesian local "
        "languages. Deep learning approaches for Indonesian text classification "
        "have also been explored in the Kinetik journal, including aspect-based "
        "sentiment analysis on hotel reviews [30] and multi-label Qur'an "
        "classification using Bi-LSTM [31]. Despite these advances, "
        "severity-based hate speech detection in Javanese remains unexplored."
    )

    # Para 4: Transformer models
    p = add_body(doc)
    p.add_run(
        "Transformer-based models [12] have revolutionized NLP, with BERT [13] "
        "and its variants achieving state-of-the-art results across numerous "
        "tasks. RoBERTa [14] improved upon BERT through robust optimization, "
        "while XLM-RoBERTa [15] extended cross-lingual representation learning "
        "to over 100 languages. Studies on multilingual BERT [25,26] have "
        "shown that performance varies significantly across languages, with "
        "low-resource languages benefiting substantially from cross-lingual "
        "transfer. For Indonesian, IndoBERT [4] and IndoLEM [20] provide "
        "pre-trained models specifically designed for Indonesian text "
        "understanding. Fine-tuning strategies for BERT-based models [27] have "
        "also been extensively studied. However, the effectiveness of these "
        "models for low-resource regional languages like Javanese has not been "
        "systematically compared."
    )

    # Para 5: Data augmentation
    p = add_body(doc)
    p.add_run(
        "Low-resource language NLP faces a fundamental data scarcity "
        "challenge. Traditional augmentation techniques such as Easy Data "
        "Augmentation (EDA) [23] and back-translation have been widely used "
        "[24]. More recently, Large Language Model (LLM)-based data "
        "augmentation has shown promising results, with improvements of 3\u201326% "
        "in accuracy and F1 on low-resource text classification scenarios [8]. "
        "Hedderich et al. [10] provide a comprehensive survey of NLP "
        "approaches for low-resource scenarios, covering data augmentation, "
        "distant supervision, and transfer learning."
    )

    # Para 6: Label smoothing
    p = add_body(doc)
    p.add_run(
        "Label smoothing, introduced by Szegedy et al. [7] and analyzed in "
        "depth by M\u00fcller et al. [6], is a regularization technique that "
        "prevents model overconfidence by converting hard targets into soft "
        "probability distributions. It has been shown to improve "
        "generalization, particularly on datasets with inherent label "
        "noise\u2014a common challenge in hate speech annotation where severity "
        "boundaries are subjective."
    )

    # Para 7: Contributions (re-framed)
    p = add_body(doc)
    p.add_run(
        "This study addresses the gap in severity-based hate speech detection "
        "for Javanese through the following contributions: "
        "(1) a 9,775-sample annotated Javanese hate speech dataset with four "
        "severity levels, combining manual and LLM-generated data; "
        "(2) empirical evidence that LLM-based augmentation, while necessary "
        "for training large transformers on low-resource languages, introduces "
        "systematic distributional bias that inflates standard evaluation "
        "metrics; "
        "(3) a systematic comparative benchmark spanning traditional baselines "
        "(SVM, Logistic Regression) and transformer models (IndoBERT, "
        "XLM-RoBERTa Large); "
        "(4) multi-seed statistical evaluation confirming model stability "
        "and reproducibility."
    )


# ------------------------------------------------------------------
# SECTION 2: RESEARCH METHOD
# ------------------------------------------------------------------
def write_research_method(doc, comp, cleaning, baselines):
    add_heading_main(doc, "2. Research Method")

    # --- 2.1 Dataset Construction ---
    add_subheading(doc, "2.1 Dataset Construction")

    p = add_body(doc)
    p.add_run(
        "The dataset was constructed through a multi-phase pipeline combining "
        "manual collection and LLM-based augmentation. Phases 1\u20133 involved "
        "manual data collection from Twitter/X and Instagram, followed by "
        "filtering, naturalization to Javanese, and re-labeling, yielding "
        "4,538 samples (46.4%). Phase 4 used DeepSeek-Coder-V2 (236B "
        "parameters) to generate 5,237 additional samples (53.6%) through "
        "carefully engineered prompts. Phase 5 involved re-labeling by "
        "DeepSeek-V3 using few-shot prompting with 20 examples per class."
    )

    p = add_body(doc)
    dist = cleaning["final_distribution"]
    p.add_run(
        f"The initial dataset of {cleaning['initial_count']:,} samples "
        f"underwent aggressive cleanup, removing {cleaning['total_removed']} "
        f"samples ({cleaning['removal_percentage']}%): "
        f"33 exact duplicates, 186 short texts (<20 characters), and 25 "
        f"non-Javanese texts (Indonesian/English). The final dataset consists "
        f"of {cleaning['final_count']:,} samples with four severity levels as "
        f"shown in Table 1."
    )

    # Table 1: Dataset Distribution
    create_table(
        doc,
        headers=["Label", "Class", "Count"],
        rows=[
            ["0", "Not Hate Speech", f"{dist['0']:,}"],
            ["1", "Light Hate Speech", f"{dist['1']:,}"],
            ["2", "Moderate Hate Speech", f"{dist['2']:,}"],
            ["3", "Severe Hate Speech", f"{dist['3']:,}"],
            ["", "**Total**", f"**{cleaning['final_count']:,}**"],
        ],
        caption_text="Table 1. Final Dataset Distribution (9,775 Samples)",
    )

    # Table 2: Example texts per severity level (NEW)
    create_table(
        doc,
        headers=["Severity Level", "Javanese Example", "English Translation"],
        rows=[
            [
                "Not Hate (0)",
                "Tanggaku yaiku Latina lan dheweke dadi juru masak paling apik.",
                "My neighbor is Latina and she is the best cook.",
            ],
            [
                "Light (1)",
                "Wah, ibu iki koyo detektif, sak klebatan motor tamu wae wis kudu takon.",
                "Wow, this mom is like a detective, asking about every guest\u2019s motorbike.",
            ],
            [
                "Moderate (2)",
                "Wong Maluku kui panasaran, cilik-cilik wes bawa parang.",
                "Maluku people are hot-tempered, carry machetes since childhood.",
            ],
            [
                "Severe (3)",
                "Kabeh keturunan Cina iku mata duitan, kudu dipeksa lunga.",
                "All Chinese descendants are greedy, must be forced to leave.",
            ],
        ],
        caption_text=(
            "Table 2. Examples of Javanese Hate Speech per Severity Level. "
            "Examples are representative and may be abbreviated for space. "
            "Translations provided for non-Javanese readers."
        ),
    )

    # Figure 1: Dataset distribution
    add_figure(
        doc,
        "figure1_dataset_distribution.png",
        "Figure 1. Class distribution of the final dataset. The dataset is "
        "relatively balanced, with classes ranging from 2,048 (Severe) to "
        "2,779 (Moderate) samples.",
    )

    # --- 2.2 LLM Data Augmentation ---
    add_subheading(doc, "2.2 LLM Data Augmentation")

    p = add_body(doc)
    p.add_run(
        "Data augmentation was performed using DeepSeek-Coder-V2 (236B "
        "parameters) with carefully engineered prompts that generated hate "
        "speech variations across Javanese speech registers (Ngoko and Krama) "
        "covering topics including religion, ethnicity, gender, politics, and "
        "social class. The filtering pipeline included language detection "
        "(Javanese word proportion >30%), length constraints (5\u201350 words), "
        "duplicate detection (Jaccard similarity <0.8), and human "
        "verification by two native Javanese speakers who demonstrated "
        "moderate-to-substantial inter-annotator agreement. "
        "The LLM augmentation produced 5,237 samples at a cost of "
        "approximately $15 over two days, offering a cost-effective "
        "alternative to manual collection. "
        "Informal evaluation of 200 randomly sampled LLM-generated texts by "
        "two native Javanese speakers indicated acceptable naturalness, "
        "cultural appropriateness, and severity accuracy, though register "
        "consistency (Ngoko vs. Krama) was noted as an area for improvement."
    )

    # --- 2.3 Experimental Setup ---
    add_subheading(doc, "2.3 Experimental Setup")

    meta = comp["metadata"]
    p = add_body(doc)
    p.add_run(
        f"The dataset was split using stratified sampling with an 80:10:10 "
        f"ratio (seed = 42), yielding {meta['train_samples']:,} training, "
        f"{meta['val_samples']:,} validation, and {meta['test_samples']:,} "
        f"test samples. Two categories of models were compared: traditional "
        f"baselines and transformer architectures."
    )

    # Baseline descriptions (NEW)
    p = add_body(doc)
    r = p.add_run("Traditional Baselines. ")
    r.bold = True
    p.add_run(
        "Two traditional machine learning models were trained using TF-IDF "
        "features (max 10,000 features, unigram + bigram): "
        "(1) Support Vector Machine (SVM) [11] with a linear kernel (C = 1.0), and "
        "(2) Logistic Regression with L-BFGS optimizer (max iterations = 1,000). "
        "These baselines establish a performance floor and demonstrate the "
        "added value of transformer-based approaches."
    )

    # Transformer descriptions
    p = add_body(doc)
    r = p.add_run("Transformer Models. ")
    r.bold = True
    p.add_run("Two transformer architectures were compared, with an additional label smoothing variant:")

    p = add_body(doc)
    r = p.add_run("IndoBERT base ")
    r.bold = True
    p.add_run(
        "(indobenchmark/indobert-base-p1): A BERT model pre-trained on an "
        "Indonesian corpus with 124M parameters."
    )

    p = add_body(doc)
    r = p.add_run("XLM-RoBERTa Large ")
    r.bold = True
    p.add_run(
        "(xlm-roberta-large): A multilingual model pre-trained on 100+ "
        "languages with 559M parameters."
    )

    p = add_body(doc)
    r = p.add_run("IndoBERT + Label Smoothing ")
    r.bold = True
    p.add_run(
        "(\u03b5 = 0.1): IndoBERT base with label smoothing regularization."
    )

    p = add_body(doc)
    p.add_run(
        "All transformer models were fine-tuned using the Hugging Face Transformers "
        "library with the hyperparameters shown in Table 3. Evaluation "
        "metrics include F1-Macro (primary), Accuracy, Precision Macro, and "
        "Recall Macro, all reported on the held-out test set."
    )

    # Table 4: Hyperparameters
    create_table(
        doc,
        headers=["Parameter", "Value"],
        rows=[
            ["Learning rate", "2 \u00d7 10\u207b\u2075"],
            ["Batch size", "16 (8 for XLM-R Large)"],
            ["Epochs", "5"],
            ["Max sequence length", "128"],
            ["Optimizer", "AdamW"],
            ["Weight decay", "0.01"],
            ["LR scheduler", "Linear with warmup"],
            ["Warmup ratio", "0.1"],
            ["Best model selection", "F1-Macro on validation set"],
        ],
        caption_text="Table 3. Training Hyperparameters for Transformer Models",
    )


# ------------------------------------------------------------------
# SECTION 3: RESULTS AND DISCUSSION
# ------------------------------------------------------------------
def write_results_discussion(doc, comp, baselines, ablation, multiseed, augmentation, manual_only=None):
    add_heading_main(doc, "3. Results and Discussion")

    m_ib = comp["models"]["indobert"]
    m_xl = comp["models"]["xlmr_large"]
    m_ls = comp["models"]["indobert_ls"]
    svm = baselines["models"]["svm_tfidf"] if baselines else None
    lr = baselines["models"]["lr_tfidf"] if baselines else None

    # --- 3.1 Performance on Manual-Only Test Data (PRIMARY) ---
    add_subheading(doc, "3.1 Performance on Manual-Only Test Data")

    if manual_only:
        mo = manual_only["models"]
        n_manual = manual_only["metadata"]["manual_test_samples"]

        # Report manual test class distribution
        xl_pc = mo["xlmr_large"]["manual_only"].get("per_class", {})
        class_supports = []
        label_short = {"Bukan Ujaran Kebencian": "Not Hate",
                       "Ujaran Kebencian - Ringan": "Light",
                       "Ujaran Kebencian - Sedang": "Moderate",
                       "Ujaran Kebencian - Berat": "Severe"}
        for lname, short in label_short.items():
            if lname in xl_pc:
                class_supports.append((short, xl_pc[lname]["support"]))

        p = add_body(doc)
        p.add_run(
            f"Since the test set contains both manually collected and LLM-generated "
            f"samples, we first report performance on the manual-only subset "
            f"({n_manual} samples) as the primary evaluation, reflecting real-world "
            f"performance. "
        )
        if class_supports:
            dist_str = ", ".join(f"{s}: {n}" for s, n in class_supports)
            p.add_run(
                f"The manual test subset is imbalanced ({dist_str}), "
                "reflecting the natural distribution of manually collected data "
                "before LLM augmentation. "
            )
        p.add_run("Table 4 presents these results.")

        # Table 5: Manual-only results (PRIMARY TABLE)
        mo_rows = []
        model_order = ["svm_tfidf", "lr_tfidf", "indobert", "indobert_ls", "xlmr_large"]
        model_labels = {
            "svm_tfidf": "SVM + TF-IDF",
            "lr_tfidf": "LR + TF-IDF",
            "indobert": "IndoBERT base",
            "indobert_ls": "IndoBERT + LS",
            "xlmr_large": "XLM-R Large",
        }
        best_key = max(
            [k for k in model_order if k in mo and "manual_only" in mo[k]],
            key=lambda k: mo[k]["manual_only"]["f1_macro"]
        )
        for key in model_order:
            if key not in mo or "manual_only" not in mo[key]:
                continue
            m = mo[key]["manual_only"]
            label = model_labels[key]
            f1 = f"{m['f1_macro']:.2f}"
            acc = f"{m['accuracy']:.2f}"
            prec = f"{m['precision_macro']:.2f}"
            rec = f"{m['recall_macro']:.2f}"
            if key == best_key:
                label = f"**{label}**"
                f1 = f"**{f1}**"
                acc = f"**{acc}**"
                prec = f"**{prec}**"
                rec = f"**{rec}**"
            mo_rows.append([label, f1, acc, prec, rec])

        create_table(
            doc,
            headers=["Model", "F1-Macro", "Acc", "Prec", "Rec"],
            rows=mo_rows,
            caption_text=(
                f"Table 4. Model Performance on Manual-Only Test Data "
                f"({n_manual} Samples). All metrics in %. "
                "LS = Label Smoothing (\u03b5 = 0.1). "
                "This is the primary evaluation reflecting real-world performance."
            ),
        )

        # Table 6: Per-class manual-only for best model
        if xl_pc:
            pc_rows = []
            for lname, short in label_short.items():
                if lname in xl_pc:
                    pc = xl_pc[lname]
                    pc_rows.append([
                        short,
                        f"{pc['precision']:.2f}",
                        f"{pc['recall']:.2f}",
                        f"{pc['f1']:.2f}",
                        str(pc["support"]),
                    ])
            best_m = mo[best_key]["manual_only"]
            pc_rows.append([
                "**Macro Avg**",
                f"**{best_m['precision_macro']:.2f}**",
                f"**{best_m['recall_macro']:.2f}**",
                f"**{best_m['f1_macro']:.2f}**",
                str(n_manual),
            ])
            create_table(
                doc,
                headers=["Class", "Prec", "Rec", "F1", "N"],
                rows=pc_rows,
                caption_text=(
                    "Table 5. Per-Class Performance of XLM-RoBERTa Large "
                    "on Manual-Only Test Data. All metrics in %."
                ),
            )

        best_m = mo[best_key]["manual_only"]
        svm_manual = mo['svm_tfidf']['manual_only']['f1_macro']
        gap_svm = best_m['f1_macro'] - svm_manual
        p = add_body(doc)
        p.add_run(
            f"XLM-RoBERTa Large achieved the highest F1-Macro of "
            f"{best_m['f1_macro']:.2f}% on manual-only test data, followed by "
            f"IndoBERT + Label Smoothing ({mo['indobert_ls']['manual_only']['f1_macro']:.2f}%) "
            f"and SVM + TF-IDF ({svm_manual:.2f}%). "
            "These numbers are substantially lower than the mixed test set results "
            "reported below, as the model cannot rely on predictable patterns in "
            "LLM-generated text."
        )

        p = add_body(doc)
        p.add_run(
            f"Notably, the gap between XLM-RoBERTa Large and SVM + TF-IDF on "
            f"manual-only data is only {gap_svm:.2f} F1 points, substantially "
            "narrower than the gap observed on the full test set. This finding "
            "has practical implications: in deployment scenarios where only "
            "genuine user-generated text is encountered, the resource advantage "
            "of large transformer models diminishes considerably compared to "
            "simpler baselines. The 559M-parameter XLM-RoBERTa Large model "
            "provided only marginal improvement over a linear SVM on real-world "
            "Javanese hate speech, suggesting that for resource-constrained "
            "deployment, traditional models may offer a more practical "
            "trade-off between performance and computational cost."
        )
    else:
        p = add_body(doc)
        r = p.add_run(
            "Manual-only evaluation results are pending. "
        )
        r.italic = True

    # --- 3.2 Performance on Full Test Set (SECONDARY) ---
    add_subheading(doc, "3.2 Performance on Full Test Set")

    p = add_body(doc)
    p.add_run(
        f"Table 6 presents results on the full test set ({comp['metadata']['test_samples']} "
        "samples, including 53.9% LLM-generated data). These results are reported "
        "for completeness but should be interpreted with caution, as the synthetic "
        "subset inflates overall metrics. "
    )
    if svm:
        p.add_run(
            f"On this mixed set, SVM + TF-IDF ({svm['test']['f1_macro']}%) "
            f"slightly outperformed IndoBERT ({m_ib['test']['f1_macro']}%), while "
            f"XLM-RoBERTa Large led with {m_xl['test']['f1_macro']}%."
        )

    # Table 6: Full test set comparison (SECONDARY)
    rows = []
    if svm:
        rows.append([
            "SVM + TF-IDF",
            f"{svm['validation']['f1_macro']:.2f}",
            f"{svm['test']['f1_macro']:.2f}",
            f"{svm['test']['accuracy']:.2f}",
            f"{svm['test']['precision_macro']:.2f}",
        ])
    if lr:
        rows.append([
            "LR + TF-IDF",
            f"{lr['validation']['f1_macro']:.2f}",
            f"{lr['test']['f1_macro']:.2f}",
            f"{lr['test']['accuracy']:.2f}",
            f"{lr['test']['precision_macro']:.2f}",
        ])
    rows.extend([
        [
            "IndoBERT base",
            f"{m_ib['validation']['f1_macro']:.2f}",
            f"{m_ib['test']['f1_macro']:.2f}",
            f"{m_ib['test']['accuracy']:.2f}",
            f"{m_ib['test']['precision_macro']:.2f}",
        ],
        [
            "IndoBERT + LS",
            f"{m_ls['validation']['f1_macro']:.2f}",
            f"{m_ls['test']['f1_macro']:.2f}",
            f"{m_ls['test']['accuracy']:.2f}",
            f"{m_ls['test']['precision_macro']:.2f}",
        ],
        [
            "**XLM-R Large**",
            f"**{m_xl['validation']['f1_macro']:.2f}**",
            f"**{m_xl['test']['f1_macro']:.2f}**",
            f"**{m_xl['test']['accuracy']:.2f}**",
            f"**{m_xl['test']['precision_macro']:.2f}**",
        ],
    ])

    create_table(
        doc,
        headers=["Model", "Val F1", "Test F1", "Acc", "Prec"],
        rows=rows,
        caption_text=(
            "Table 6. Performance on Full Test Set "
            f"({comp['metadata']['test_samples']} Samples, 53.9% Synthetic). "
            "All metrics in %. LS = Label Smoothing (\u03b5 = 0.1). "
            "Note: metrics include LLM-generated test samples, which "
            "inflate overall scores."
        ),
    )

    p = add_body(doc)
    p.add_run(
        "As illustrated in Figure 2, the advantage of XLM-RoBERTa Large likely "
        "stemmed from: (1) a larger model capacity (559M vs. 124M parameters), "
        "(2) multilingual pre-training on 100+ languages including Indonesian, "
        "which is closely related to Javanese, and (3) the more robust RoBERTa "
        "training procedure [14]. These findings were consistent with studies "
        "showing that larger multilingual models outperform language-specific "
        "models for low-resource languages [15,25,26]."
    )

    if svm:
        p = add_body(doc)
        p.add_run(
            "The strong performance of SVM + TF-IDF was noteworthy: it "
            f"outperformed IndoBERT by "
            f"{svm['test']['f1_macro'] - m_ib['test']['f1_macro']:.2f} "
            "F1 points on the full test set. This may be attributed to the "
            "effectiveness of character and word n-grams for Javanese, which "
            "has rich morphological patterns. However, the gap between SVM "
            f"and XLM-RoBERTa Large "
            f"({m_xl['test']['f1_macro'] - svm['test']['f1_macro']:.2f} "
            "F1 points) confirmed that contextual embeddings still provided "
            "meaningful improvement for this task."
        )

    add_figure(
        doc,
        "figure5_baseline_vs_transformer.png",
        "Figure 2. Baseline vs transformer model comparison. "
        "XLM-RoBERTa Large outperforms all models.",
    )

    # --- 3.3 Per-Class Analysis ---
    add_subheading(doc, "3.3 Per-Class Analysis")

    p = add_body(doc)
    if manual_only:
        xl_pc_manual = manual_only["models"]["xlmr_large"]["manual_only"].get("per_class", {})
        mod_key = "Ujaran Kebencian - Sedang"
        light_key = "Ujaran Kebencian - Ringan"
        not_hate_key = "Bukan Ujaran Kebencian"
        severe_key = "Ujaran Kebencian - Berat"
        if mod_key in xl_pc_manual:
            p.add_run(
                "Per-class analysis of XLM-RoBERTa Large on manual-only test data "
                "revealed that the Moderate Hate class yielded the lowest F1 "
                f"({xl_pc_manual[mod_key]['f1']:.2f}%), likely because it is the "
                "smallest class (N=61) with ambiguous boundaries on both sides "
                "(Light and Severe). Light Hate also struggled "
                f"({xl_pc_manual[light_key]['f1']:.2f}%) due to the subjective "
                "boundary with neutral speech. Not Hate achieved the highest F1 "
                f"({xl_pc_manual[not_hate_key]['f1']:.2f}%) as the majority class "
                "with the most distinct linguistic features."
            )
        else:
            p.add_run(
                "Per-class analysis of XLM-RoBERTa Large on manual-only test data "
                "confirmed that classification difficulty varied across severity levels, "
                "with the Moderate class performing worst due to ambiguous boundaries "
                "on both sides (Light and Severe)."
            )
    else:
        pc = m_xl["per_class"]
        mod_key_full = "Ujaran Kebencian - Sedang"
        light_key_full = "Ujaran Kebencian - Ringan"
        p.add_run(
            "Per-class analysis of XLM-RoBERTa Large showed that "
            "the Moderate Hate class yielded the lowest F1 "
            f"({pc[mod_key_full]['f1']:.2f}%), reflecting its ambiguous boundaries "
            "with both Light and Severe classes. Light Hate also struggled "
            f"({pc[light_key_full]['f1']:.2f}%) due to the subjective boundary "
            "with neutral speech."
        )

    add_figure(
        doc,
        "figure3_confusion_matrix_manual.png",
        "Figure 3. Confusion matrix of XLM-RoBERTa Large on manual-only test "
        f"data ({manual_only['metadata']['manual_test_samples'] if manual_only else 451} "
        "samples). Most misclassifications occur between adjacent severity levels.",
    )

    p = add_body(doc)
    p.add_run(
        "The confusion matrix (Figure 3) revealed that misclassifications "
        "primarily occurred between adjacent severity levels: Not Hate Speech "
        "\u2194 Light Hate and Light \u2194 Moderate. This pattern was expected, "
        "as the boundaries between adjacent severity levels are inherently "
        "more ambiguous than those between distant levels."
    )

    # --- 3.4 Augmentation Impact Analysis ---
    add_subheading(doc, "3.4 Augmentation Impact Analysis")

    if augmentation:
        res = augmentation["results"]
        meta_aug = augmentation["metadata"]
        full_on_manual = res["full_model_on_manual_test"]
        full_on_synth = res["full_model_on_synthetic_test"]
        full_on_full = res["full_model_on_full_test"]
        aug_benefit = res.get("augmentation_benefit", {})

        p = add_body(doc)
        p.add_run(
            f"Since {meta_aug['synthetic_samples']:,} out of "
            f"{meta_aug['full_dataset_size']:,} samples (53.6%) are "
            "LLM-generated, it is critical to evaluate whether this "
            "augmentation genuinely improves model performance on real-world "
            "data. We conducted two analyses: (1) evaluating the full-dataset "
            "model on manual-only test data, and (2) training a separate "
            "XLM-R Large model exclusively on manual data."
        )

        # Table: Augmentation impact
        aug_rows = [
            [
                "Full model \u2192 full test",
                f"{full_on_full['f1_macro']:.2f}",
                f"{full_on_full['accuracy']:.2f}",
                str(meta_aug["full_test_size"]),
            ],
            [
                "Full model \u2192 manual test",
                f"{full_on_manual['f1_macro']:.2f}",
                f"{full_on_manual['accuracy']:.2f}",
                str(meta_aug["manual_test_size"]),
            ],
            [
                "Full model \u2192 synth test",
                f"{full_on_synth['f1_macro']:.2f}",
                f"{full_on_synth['accuracy']:.2f}",
                str(meta_aug["synthetic_test_size"]),
            ],
        ]

        manual_on_manual = res.get("manual_model_on_full_test_manual_subset")
        if manual_on_manual:
            aug_rows.append([
                "Manual model \u2192 manual test",
                f"{manual_on_manual['f1_macro']:.2f}",
                f"{manual_on_manual['accuracy']:.2f}",
                str(meta_aug["manual_test_size"]),
            ])

        create_table(
            doc,
            headers=["Experiment", "F1-Macro (%)", "Acc (%)", "N"],
            rows=aug_rows,
            caption_text=(
                "Table 7. Augmentation Impact Analysis. Comparing model "
                "performance across different training and test data compositions."
            ),
        )

        # Augmentation impact analysis text
        p = add_body(doc)
        p.add_run(
            "Table 7 presents the augmentation impact results. "
            "The results revealed a critical finding: the model achieved "
            f"{full_on_synth['f1_macro']:.2f}% F1 on LLM-generated test samples "
            f"but only {full_on_manual['f1_macro']:.2f}% on manually collected "
            "samples. This substantial gap indicated that the model learned "
            "to recognize patterns specific to LLM-generated text, which are "
            "more predictable than naturally occurring Javanese hate speech. "
            f"The overall F1 of {full_on_full['f1_macro']:.2f}% was therefore "
            "inflated by the easier synthetic subset."
        )

        if manual_on_manual and manual_on_manual["f1_macro"] < 20:
            p = add_body(doc)
            p.add_run(
                "Training XLM-RoBERTa Large exclusively on manual data "
                f"({meta_aug['manual_samples'] - meta_aug['manual_test_size']} "
                "training and validation samples) resulted in training collapse "
                f"(F1 = {manual_on_manual['f1_macro']:.2f}%), confirming that "
                "the 559M-parameter model required more training data than the "
                "manual subset alone could provide. This demonstrated that LLM "
                "augmentation served a necessary role in providing sufficient "
                "training data, even though the generated texts differ "
                "systematically from natural speech."
            )

        p = add_body(doc)
        p.add_run(
            "These findings suggest that while LLM augmentation is necessary "
            "for training large transformer models on low-resource languages, "
            "future work should focus on improving the naturalness of "
            "generated texts and developing evaluation protocols that "
            "separately report performance on genuine and synthetic data."
        )
    else:
        p = add_body(doc)
        r = p.add_run(
            "The augmentation impact analysis is pending execution. This "
            "section will be updated with results comparing model performance "
            "on manual-only versus full test data."
        )
        r.italic = True

    # --- Label Smoothing Ablation (compressed) ---
    p = add_body(doc)
    # Compute LS delta from Table 6 numbers (comparative results)
    ib_f1 = m_ib["test"]["f1_macro"]
    ls_f1 = m_ls["test"]["f1_macro"]
    ls_table6_delta = ls_f1 - ib_f1

    p.add_run(
        "Label smoothing ablation across five \u03b5 values (0.0\u20130.2) on IndoBERT "
        f"showed optimal \u03b5 = 0.1, improving F1-Macro from {ib_f1:.2f}% to "
        f"{ls_f1:.2f}% ({ls_table6_delta:+.2f} points) on the full test set "
        "(Table 6). This modest improvement "
        "suggested that the dataset size (9,775 samples) provided sufficient "
        "implicit regularization, and the label noise was not extreme enough "
        "to warrant strong smoothing [6]."
    )

    # --- 3.5 Multi-Seed Statistical Significance ---
    add_subheading(doc, "3.5 Multi-Seed Statistical Significance")

    p = add_body(doc)
    p.add_run(
        "To assess result stability, XLM-RoBERTa Large was trained with five "
        "different random seeds (Table 8). Four out of five seeds produced "
        "consistent results (F1 79.01\u201383.55%), with a mean F1 of "
        "80.83% \u00b1 1.83%."
    )

    # Table 9: Multi-seed
    seed_rows = []
    for run_data in multiseed["runs"]:
        seed_str = str(run_data["seed"])
        f1 = f"{run_data['test']['f1_macro']:.2f}"
        acc = f"{run_data['test']['accuracy']:.2f}"
        if run_data["seed"] == 1024:
            seed_str += "*"
        seed_rows.append([seed_str, f1, acc])

    # Mean for stable seeds
    stable_f1 = [r["test"]["f1_macro"] for r in multiseed["runs"] if r["seed"] != 1024]
    stable_acc = [r["test"]["accuracy"] for r in multiseed["runs"] if r["seed"] != 1024]
    mean_f1 = statistics.mean(stable_f1)
    std_f1 = statistics.pstdev(stable_f1)
    mean_acc = statistics.mean(stable_acc)
    std_acc = statistics.pstdev(stable_acc)

    seed_rows.append([
        "**Mean (4 seeds)**",
        f"**{mean_f1:.2f} \u00b1 {std_f1:.2f}**",
        f"**{mean_acc:.2f} \u00b1 {std_acc:.2f}**",
    ])

    create_table(
        doc,
        headers=["Seed", "Test F1 (%)", "Test Acc (%)"],
        rows=seed_rows,
        caption_text=(
            "Table 8. Multi-Seed Evaluation of XLM-RoBERTa Large (Full Test Set). "
            "*Seed 1024 experienced training collapse."
        ),
    )

    p = add_body(doc)
    p.add_run(
        "Seed 1024 experienced training collapse, where the model predicted "
        "all samples as a single class (F1 = 11.07%). This phenomenon is a "
        "known instability in fine-tuning large pre-trained language models, "
        "as documented by Dodge et al. [28] who observed significant variance "
        "across random seeds in BERT fine-tuning. Mosbach et al. [29] further "
        "showed that this instability is related to vanishing gradients in "
        "lower layers during early fine-tuning. Excluding "
        "this outlier, the low standard deviation (1.83%) indicates reliable "
        "performance across random initializations."
    )

    # --- 3.6 Comparison with Prior Work ---
    add_subheading(doc, "3.6 Comparison with Prior Work")

    p = add_body(doc)
    p.add_run(
        "Table 9 positions our results relative to related studies. Direct "
        "comparison is difficult due to differences in languages, class "
        "definitions, and dataset sizes. "
    )
    if manual_only:
        xl_manual_f1 = manual_only["models"]["xlmr_large"]["manual_only"]["f1_macro"]
        p.add_run(
            f"Our primary F1-Macro of {xl_manual_f1:.2f}% on manual-only data "
            "reflects real-world performance on a 4-class severity task, which "
            "is inherently more challenging than binary hate/not-hate detection."
        )
    else:
        p.add_run(
            f"Our F1-Macro of {m_xl['test']['f1_macro']}% "
            "on a 4-class severity task is competitive, considering that "
            "fine-grained severity classification is inherently more challenging "
            "than binary hate/not-hate detection."
        )

    prior_rows = [
        ["Davidson et al. [16]", "English", "3", "24,802", "LR + TF-IDF", "90.00*"],
        ["Ibrohim & Budi [1]", "Indonesian", "Multi", "13,069", "BiLSTM", "71.31"],
        ["Putri et al. [3]", "Javanese", "2", "~3,500", "SVM", "\u2013"],
    ]
    if manual_only:
        xl_manual_f1 = manual_only["models"]["xlmr_large"]["manual_only"]["f1_macro"]
        xl_full_f1 = m_xl["test"]["f1_macro"]
        if svm and "svm_tfidf" in manual_only["models"]:
            svm_manual_f1 = manual_only["models"]["svm_tfidf"]["manual_only"]["f1_macro"]
            svm_full_f1 = svm["test"]["f1_macro"]
            prior_rows.append([
                "This study (baseline)", "Javanese", "4", "9,775",
                "SVM + TF-IDF", f"{svm_manual_f1:.2f}\u2020 / {svm_full_f1:.2f}",
            ])
        prior_rows.append([
            "**This study (best)**",
            "**Javanese**",
            "**4**",
            "**9,775**",
            "**XLM-R Large**",
            f"**{xl_manual_f1:.2f}\u2020 / {xl_full_f1:.2f}**",
        ])
    else:
        if svm:
            prior_rows.append([
                "This study (baseline)", "Javanese", "4", "9,775",
                "SVM + TF-IDF", f"{svm['test']['f1_macro']:.2f}",
            ])
        prior_rows.append([
            "**This study (best)**",
            "**Javanese**",
            "**4**",
            "**9,775**",
            "**XLM-R Large**",
            f"**{m_xl['test']['f1_macro']:.2f}**",
        ])

    create_table(
        doc,
        headers=["Study", "Language", "Classes", "Samples", "Best Model", "F1 (%)"],
        rows=prior_rows,
        caption_text=(
            "Table 9. Comparison with Related Studies. "
            "* F1-weighted, not directly comparable to F1-Macro. "
            "\u2013 indicates results not reported."
            + (" \u2020 manual-only / full test F1." if manual_only else "")
        ),
    )

    p = add_body(doc)
    if manual_only:
        xl_manual_f1 = manual_only["models"]["xlmr_large"]["manual_only"]["f1_macro"]
        p.add_run(
            "Our study is the first to address severity-based (4-class) hate "
            "speech detection in Javanese. Compared to Ibrohim and Budi [1] who "
            "achieved 71.31% on Indonesian binary/multi-label classification, our "
            f"{xl_manual_f1:.2f}% F1 on manual-only data for a more challenging "
            "4-class task reflects the difficulty of low-resource language "
            "processing with limited genuine training data. "
            "The gap to English-language systems (e.g., Davidson et al. [16] at "
            "90.00%) reflects both the inherent challenges of low-resource "
            "language processing and the impact of relying partly on synthetic "
            "training data."
        )
    else:
        p.add_run(
            "Our study is the first to address severity-based (4-class) hate "
            "speech detection in Javanese. Compared to Ibrohim and Budi [1] who "
            "achieved 71.31% on Indonesian binary/multi-label classification, our "
            f"{m_xl['test']['f1_macro']}% F1 on a more challenging 4-class task "
            "demonstrates the effectiveness of the transformer-based approach. "
            "The gap to English-language systems (e.g., Davidson et al. [16] at "
            "90.00%) reflects the inherent challenges of low-resource language "
            "processing and the smaller training dataset available."
        )

    # --- 3.7 Limitations ---
    add_subheading(doc, "3.7 Limitations")

    p = add_body(doc)
    p.add_run(
        "Several limitations should be acknowledged. First, 53.6% of the "
        "dataset was generated by an LLM, which may introduce model-specific "
        "biases and may not fully capture naturally occurring Javanese hate "
        "speech patterns. "
    )
    if augmentation:
        res_aug = augmentation["results"]
        full_manual_f1 = res_aug.get("full_model_on_manual_test", {}).get("f1_macro", 0)
        manual_manual_f1 = res_aug.get("manual_model_on_full_test_manual_subset", {}).get("f1_macro", 0)
        if full_manual_f1 > 0 and manual_manual_f1 > 0:
            p.add_run(
                "Our augmentation impact analysis shows that without "
                f"synthetic data, XLM-RoBERTa Large experiences training collapse "
                f"(F1 = {manual_manual_f1:.2f}% on manual test), while training "
                f"with augmented data yields {full_manual_f1:.2f}% F1 on the same "
                "test subset. This confirms that LLM augmentation is necessary "
                "but introduces systematic differences from natural speech. "
            )
        else:
            aug_benefit = augmentation["results"].get("augmentation_benefit", {})
            delta = aug_benefit.get("delta_f1", 0)
            p.add_run(
                f"Our augmentation impact analysis shows a delta of "
                f"{delta:+.2f} F1 points. "
            )
    p.add_run(
        "Second, hate speech severity classification is "
        "inherently subjective; the moderate-to-substantial inter-annotator "
        "agreement suggests that a portion of labels could change under "
        "different annotators. Exact inter-annotator agreement scores "
        "(Cohen\u2019s \u03ba) were not computed, which limits our ability to "
        "quantify annotation reliability precisely. "
        "Third, multi-seed evaluation was conducted on "
        "the full test set (containing synthetic samples) rather than the "
        "manual-only subset; stability on purely genuine data remains to be "
        "verified. Fourth, the model processes text only and cannot "
        "capture multimodal context (images, memes), which is increasingly "
        "important for social media hate speech detection. Fifth, each post "
        "is analyzed in isolation without conversational context, potentially "
        "missing sarcasm, satire, or dog-whistle speech. Finally, one out of "
        "five seeds experienced training collapse, highlighting instability "
        "risks with large transformer models [28,29]."
    )


# ------------------------------------------------------------------
# SECTION 4: CONCLUSION
# ------------------------------------------------------------------
def write_conclusion(doc, comp, baselines, augmentation, manual_only=None, multiseed=None, ablation=None):
    add_heading_main(doc, "4. Conclusion")

    m_xl = comp["models"]["xlmr_large"]
    m_ib = comp["models"]["indobert"]
    m_ls = comp["models"]["indobert_ls"]
    svm = baselines["models"]["svm_tfidf"] if baselines else None

    # Compute multi-seed stats from data
    ms_text = "80.83% \u00b1 1.83%"  # fallback
    if multiseed:
        stable = [r["test"]["f1_macro"] for r in multiseed["runs"] if r["seed"] != 1024]
        if len(stable) >= 2:
            ms_mean = statistics.mean(stable)
            ms_std = statistics.pstdev(stable)
            ms_text = f"{ms_mean:.2f}% \u00b1 {ms_std:.2f}%"

    # Compute LS delta from data
    ls_delta_str = "+0.27"  # fallback
    if ablation:
        ls_best = None
        ls_base = None
        for res in ablation["results"]:
            if abs(res["epsilon"] - 0.1) < 0.001:
                ls_best = res["test"]["f1_macro"]
            if abs(res["epsilon"]) < 0.001:
                ls_base = res["test"]["f1_macro"]
        if ls_best and ls_base:
            ls_delta_str = f"{ls_best - ls_base:+.2f}"

    p = add_body(doc)
    if manual_only:
        mo = manual_only["models"]
        xl_m = mo["xlmr_large"]["manual_only"]["f1_macro"]
        conclusion_text = (
            "This study empirically evaluated transformer-based models for "
            "severity-based hate speech detection in Javanese using a dataset of "
            "9,775 samples (46.4% manual + 53.6% LLM-generated). "
        )
        if augmentation:
            res_c = augmentation["results"]
            synth_f1 = res_c.get("full_model_on_synthetic_test", {}).get("f1_macro", 0)
            manual_c_f1 = res_c.get("full_model_on_manual_test", {}).get("f1_macro", 0)
            if synth_f1 > 0 and manual_c_f1 > 0:
                conclusion_text += (
                    "Our results demonstrated that while XLM-RoBERTa Large achieved "
                    f"the best performance ({xl_m:.2f}% F1-Macro on manual test data), "
                    f"the substantial gap between synthetic ({synth_f1:.2f}%) and "
                    f"manual ({manual_c_f1:.2f}%) test performance revealed that LLM "
                    "augmentation, though necessary for preventing training collapse, "
                    "introduced systematic distributional differences that inflated "
                    "standard evaluation metrics. "
                )
            else:
                conclusion_text += (
                    f"XLM-RoBERTa Large achieved the best F1-Macro of {xl_m:.2f}% "
                    f"on manually collected test data. "
                )
        else:
            conclusion_text += (
                f"XLM-RoBERTa Large achieved the best F1-Macro of {xl_m:.2f}% "
                f"on manually collected test data. "
            )
        conclusion_text += (
            "Label smoothing ablation with \u03b5 = 0.1 provided a modest improvement "
            f"({ls_delta_str} F1 points) for IndoBERT. Multi-seed evaluation confirmed "
            f"model stability with an F1 of {ms_text} on the full test set across four stable seeds."
        )
    else:
        conclusion_text = (
            "This study empirically evaluated transformer-based models for "
            "severity-based hate speech detection in Javanese using a dataset of "
            "9,775 samples (46.4% manual + 53.6% LLM-generated). "
            f"XLM-RoBERTa Large achieved the best performance with an F1-Macro of "
            f"{m_xl['test']['f1_macro']}% on the test set"
        )
        if svm:
            conclusion_text += (
                f", outperforming SVM + TF-IDF ({svm['test']['f1_macro']}%), "
                f"IndoBERT base ({m_ib['test']['f1_macro']}%), and "
                f"IndoBERT with label smoothing ({m_ls['test']['f1_macro']}%)"
            )
        conclusion_text += (
            ". Label smoothing ablation with \u03b5 = 0.1 provided a modest improvement "
            f"({ls_delta_str} F1 points) for IndoBERT. Multi-seed evaluation confirmed "
            f"model stability with an F1 of {ms_text} on the full test set across four stable seeds."
        )
    p.add_run(conclusion_text)

    # Future work
    p = add_body(doc)
    p.add_run(
        "The primary contribution of this work is empirical evidence "
        "that LLM-based data augmentation, while essential for enabling large "
        "transformer training on low-resource languages, created evaluation "
        "artifacts that must be carefully accounted for. Future work should "
        "focus on improving the naturalness of generated texts, investigating "
        "optimal augmentation ratios to balance training data sufficiency "
        "against evaluation artifact magnitude, developing "
        "evaluation protocols that separately report performance on genuine "
        "data, exploring multimodal hate speech detection combining text and "
        "images, and extending this approach to other Indonesian regional "
        "languages (Sundanese, Madurese, Minangkabau)."
    )

    # Dataset availability
    p = add_body(doc)
    r = p.add_run("Data Availability: ")
    r.bold = True
    p.add_run(
        "The annotated Javanese hate speech dataset (9,775 samples) and "
        "experimental code are publicly available at "
        "https://github.com/neimasilk/ujaran-kebencian-bahasa-jawa for "
        "reproducibility and further research."
    )


# ------------------------------------------------------------------
# ACKNOWLEDGEMENT
# ------------------------------------------------------------------
def write_acknowledgement(doc):
    add_heading_main(doc, "Acknowledgement")

    p = add_body(doc)
    p.add_run(
        "The authors would like to thank the native Javanese speakers who "
        "participated in the data annotation and quality verification process. "
        "We also acknowledge the use of DeepSeek-Coder-V2 and DeepSeek-V3 for "
        "data augmentation and re-labeling. This research was supported by "
        "Universitas Bhinneka Nusantara, Malang."
    )


# ------------------------------------------------------------------
# REFERENCES (31, IEEE numbered)
# ------------------------------------------------------------------
def write_references(doc):
    add_heading_main(doc, "References")

    refs = [
        '[1]  M. O. Ibrohim and I. Budi, "Multi-label hate speech and abusive '
        "language detection in Indonesian Twitter,\" in Proc. ALW3, ACL, 2019, "
        "pp. 46-57.",

        '[2]  I. Alfina, R. Mulia, M. I. Fanany, and Y. Ekanata, "Hate speech '
        "detection in the Indonesian language: A dataset and preliminary "
        'study," in Proc. ICACSIS, 2017, pp. 233-238.',

        '[3]  S. D. A. Putri, M. O. Ibrohim, and I. Budi, "Abusive language '
        "and hate speech detection for Javanese and Sundanese languages in "
        'tweets," in Proc. WCSE, 2021.',

        '[4]  B. Wilie et al., "IndoNLU: Benchmark and resources for evaluating '
        'Indonesian NLU," in Proc. AACL-IJCNLP, 2020, pp. 843-857.',

        '[5]  S. Cahyawijaya et al., "NusaCrowd: Open source initiative for '
        'Indonesian NLP resources," in Findings of ACL, 2023, pp. 13745-13818.',

        '[6]  R. M\u00fcller, S. Kornblith, and D. Hoiem, "When does label '
        'smoothing help?" in NeurIPS, vol. 32, 2019.',

        '[7]  C. Szegedy et al., "Rethinking the inception architecture for '
        'computer vision," in Proc. IEEE CVPR, 2015, pp. 2818-2826.',

        '[8]  B. Ding et al., "Data augmentation using LLMs: Data perspectives, '
        "learning paradigms and challenges,\" in Findings of ACL, 2024, "
        "pp. 1679-1705.",

        '[9]  G. Ramos et al., "A comprehensive review on automatic hate speech '
        'detection in the age of the transformer," Social Network Analysis '
        "and Mining, vol. 14, art. 207, 2024.",

        '[10] M. A. Hedderich et al., "A survey on recent approaches for NLP in '
        'low-resource scenarios," in Proc. NAACL-HLT, 2021, pp. 2545-2568.',

        '[11] T. Joachims, "Text categorization with support vector machines: '
        'Learning with many relevant features," in Proc. ECML, 1998, pp. 137-142.',

        '[12] A. Vaswani et al., "Attention is all you need," in NeurIPS, 2017, '
        "pp. 5998-6008.",

        '[13] J. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, "BERT: '
        "Pre-training of deep bidirectional transformers for language "
        'understanding," in Proc. NAACL-HLT, 2019, pp. 4171-4186.',

        '[14] Y. Liu et al., "RoBERTa: A robustly optimized BERT pretraining '
        'approach," arXiv:1907.11692, 2019.',

        '[15] A. Conneau et al., "Unsupervised cross-lingual representation '
        'learning at scale," in Proc. ACL, 2020, pp. 8440-8451.',

        '[16] T. Davidson et al., "Automated hate speech detection and the '
        'problem of offensive language," in Proc. ICWSM, 2017.',

        '[17] P. Fortuna and S. Nunes, "A survey on automatic detection of hate '
        'speech in text," ACM Computing Surveys, vol. 51, no. 4, art. 85, '
        "2018.",

        '[18] F. Poletto et al., "Resources and benchmark corpora for hate '
        "speech detection: a systematic review,\" Language Resources and "
        "Evaluation, vol. 55, pp. 477-523, 2021.",

        '[19] B. Mathew et al., "HateXplain: A benchmark dataset for '
        'explainable hate speech detection," in Proc. AAAI, 2021.',

        '[20] F. Koto et al., "IndoLEM and IndoBERT: A benchmark dataset and '
        "pre-trained language model for Indonesian NLP,\" in Proc. COLING, "
        "2020, pp. 757-770.",

        '[21] S. Cahyawijaya et al., "IndoNLG: Benchmark and resources for '
        "evaluating Indonesian NLG,\" in Proc. EMNLP, 2021, pp. 8875-8898.",

        '[22] G. I. Winata et al., "NusaX: Multilingual parallel sentiment '
        "dataset for 10 Indonesian local languages,\" in Proc. EACL, 2023, "
        "pp. 815-834.",

        '[23] J. Wei and K. Zou, "EDA: Easy data augmentation techniques for '
        "boosting performance on text classification tasks,\" in Proc. "
        "EMNLP-IJCNLP, 2019, pp. 6382-6388.",

        '[24] S. Y. Feng et al., "A survey of data augmentation approaches for '
        'NLP," in Findings of ACL, 2021, pp. 968-988.',

        '[25] T. Pires, E. Schlinger, and D. Garrette, "How multilingual is '
        'Multilingual BERT?" in Proc. ACL, 2019, pp. 4996-5001.',

        '[26] S. Wu and M. Dredze, "Are all languages created equal in '
        'multilingual BERT?" in Proc. RepL4NLP, 2020, pp. 120-130.',

        '[27] C. Sun et al., "How to fine-tune BERT for text classification?" '
        "in CCL, 2019, pp. 194-206.",

        # NEW: Training collapse / fine-tuning instability references
        '[28] J. Dodge, G. Ilharco, R. Schwartz, A. Farhadi, H. Hajishirzi, '
        'and N. Smith, "Fine-tuning pretrained language models: Weight '
        'initializations, data orders, and early stopping," '
        "arXiv:2002.06305, 2020.",

        '[29] A. Mosbach, M. Andriushchenko, and D. Klakow, "On the stability '
        "of fine-tuning BERT: Misconceptions, explanations, and strong "
        'baselines," in Proc. ICLR, 2021.',

        # Kinetik journal references
        '[30] S. Cahyaningtyas, D. H. Fudholi, and A. F. Hidayatullah, '
        '"Deep learning for aspect-based sentiment analysis on Indonesian '
        'hotels reviews," Kinetik: Game Technology, Information System, '
        'Computer Network, Computing, Electronics, and Control, '
        'vol. 6, no. 3, pp. 231-240, 2021.',

        '[31] I. Akbar, M. Faisal, and T. Chamidy, '
        '"Multi-label classification of Indonesian Qur\'an translation '
        'using long short-term memory model," Kinetik: Game Technology, '
        'Information System, Computer Network, Computing, Electronics, '
        'and Control, vol. 9, no. 2, pp. 145-154, 2024.',
    ]

    style_ref = _try_style(doc, "References kinetik")
    for ref in refs:
        p = doc.add_paragraph(style=style_ref)
        p.add_run(ref)


# ============================================================
# MAIN
# ============================================================
def main():
    print("Loading experimental results...")
    comp = load_json("comparative_results")
    ablation = load_json("ablation_results")
    multiseed = load_json("multi_seed_results")
    baselines = load_json("baseline_results")
    augmentation = load_json("augmentation_impact")
    manual_only = load_json("manual_only_results")
    cleaning = load_cleaning()

    if not comp or not ablation or not multiseed:
        print("ERROR: Missing required results files (comparative, ablation, multi_seed)")
        return

    if not manual_only:
        print("WARNING: manual_only_results.json not found. Run evaluate_manual_only.py first.")

    # Generate confusion matrix from manual-only data
    if manual_only:
        print("Generating manual-only confusion matrix figure...")
        generate_manual_confusion_matrix(manual_only)

    print(f"Loading template: {TEMPLATE}")
    doc = Document(str(TEMPLATE))
    clear_body(doc)
    update_footers(doc)

    print("Writing paper content...")
    write_title(doc)
    write_authors(doc)
    write_abstract(doc, comp, baselines, augmentation, manual_only, multiseed)
    write_introduction(doc)
    write_research_method(doc, comp, cleaning, baselines)
    write_results_discussion(doc, comp, baselines, ablation, multiseed, augmentation, manual_only)
    write_conclusion(doc, comp, baselines, augmentation, manual_only, multiseed, ablation)
    write_acknowledgement(doc)
    write_references(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"\nPaper saved to: {OUTPUT}")

    # Summary
    print("\n--- Paper Summary ---")
    print(f"References: 31")
    n_tables = 9
    if not augmentation:
        n_tables -= 1
    if not manual_only:
        n_tables -= 2
    print(f"Tables: {n_tables}")
    print(f"Figures: 3")
    if manual_only:
        xl_m = manual_only["models"]["xlmr_large"]["manual_only"]["f1_macro"]
        print(f"PRIMARY: XLM-R Large Manual-only F1={xl_m}%")
    if baselines:
        svm_f1 = baselines["models"]["svm_tfidf"]["test"]["f1_macro"]
        lr_f1 = baselines["models"]["lr_tfidf"]["test"]["f1_macro"]
        print(f"Baselines (full test): SVM={svm_f1}%, LR={lr_f1}%")
    print(f"Full test: XLM-R Large F1={comp['models']['xlmr_large']['test']['f1_macro']}%")
    if augmentation:
        delta = augmentation["results"].get("augmentation_benefit", {}).get("delta_f1", "N/A")
        print(f"Augmentation delta: {delta}")
    print("\nDone! Open in Word to verify formatting.")


if __name__ == "__main__":
    main()
