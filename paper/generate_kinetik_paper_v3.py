#!/usr/bin/env python3
"""
Generate Kinetik (SINTA 2) paper DOCX — Version 3 (post-rejection rewrite).

Major changes from v2:
- New title focused on evaluation reliability
- Explicit GAP analysis in introduction
- Research questions (RQ1-RQ3)
- Comparison table with prior work (Table 1)
- Augmentation ratio study results (Table 5, Figure 2)
- Data quality analysis (Table 6, Figure 4)
- Error analysis section
- Two-stage training results (if available)
- Expanded references (~35)

Usage:
    python paper/generate_kinetik_paper_v3.py
"""

import json
import statistics
from pathlib import Path
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# PATHS
# ============================================================
BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE = BASE_DIR / "Template Kinetik Mendeley.docx"
OUTPUT = BASE_DIR / "paper" / "paper_kinetik_v3.docx"
RESULTS = BASE_DIR / "results"
FIGURES = BASE_DIR / "paper" / "figures"

# ============================================================
# DATA LOADING
# ============================================================
def load_json(name):
    path = RESULTS / f"{name}.json"
    if not path.exists():
        print(f"  [WARN] Not found: {path}")
        return None
    with open(path, encoding="utf-8") as f:
        return json.load(f)

def load_cleaning():
    path = BASE_DIR / "data" / "cleaned" / "cleaning_report.json"
    if not path.exists():
        return None
    with open(path, encoding="utf-8") as f:
        return json.load(f)

# ============================================================
# HELPER FUNCTIONS (reused from v2)
# ============================================================
def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

def update_footers(doc):
    from xml.sax.saxutils import escape as xml_escape
    cite_text = (
        "Cite: Amien, M., Sijabat, D. R., & Kanthi, Y. A. (2026). "
        "Evaluating the Reliability of LLM-Augmented Training Data for "
        "Severity-Based Hate Speech Detection in Low-Resource Javanese. "
        "Kinetik: Game Technology, Information System, Computer Network, "
        "Computing, Electronics, and Control."
    )
    cite_xml = xml_escape(cite_text)
    def _make_footer_p():
        return parse_xml(
            f'<w:p {nsdecls("w")}><w:pPr><w:pStyle w:val="Footer"/>'
            f'<w:pBdr><w:top w:val="single" w:sz="4" w:space="1" w:color="auto"/></w:pBdr>'
            f'<w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
            f'<w:i/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr></w:pPr>'
            f'<w:r><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
            f'<w:i/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
            f'<w:t xml:space="preserve">{cite_xml}</w:t></w:r></w:p>'
        )
    for section in doc.sections:
        footer = section.footer
        if not footer.is_linked_to_previous:
            ft_el = footer._element
            for child in list(ft_el):
                ft_el.remove(child)
            ft_el.append(_make_footer_p())
        fp_footer = section.first_page_footer
        if not fp_footer.is_linked_to_previous:
            fp_el = fp_footer._element
            for child in list(fp_el):
                fp_el.remove(child)
            fp_el.append(_make_footer_p())

def _try_style(doc, style_name, fallback="Normal"):
    try:
        doc.styles[style_name]
        return style_name
    except KeyError:
        return fallback

def add_body(doc, text="", bold=False, italic=False):
    style = _try_style(doc, "Body kinetik")
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    return p

def add_heading_main(doc, text):
    style = _try_style(doc, "Sub title 1 kinetik")
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p

def add_subheading(doc, text):
    style = _try_style(doc, "Body kinetik")
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = True
    p.paragraph_format.space_before = Pt(6)
    return p

def add_caption(doc, text):
    style = _try_style(doc, "Figure table kinetik")
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    return p

def add_figure(doc, filename, caption_text, width=Inches(5.0)):
    path = FIGURES / filename
    if not path.exists():
        add_body(doc, f"[Figure not found: {filename}]", italic=True)
        return
    doc.add_picture(str(path), width=width)
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption_text)

def fmt_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        "</w:tblBorders>"
    )
    tblPr.append(borders)

def set_header_shading(table, n_cols, color="D9E2F3"):
    for i in range(n_cols):
        cell = table.cell(0, i)
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
        )
        cell._element.get_or_add_tcPr().append(shading)

def create_table(doc, headers, rows, caption_text):
    add_caption(doc, caption_text)
    n_cols = len(headers)
    n_rows = len(rows) + 1
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        fmt_cell(table.cell(0, i), h, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            b = isinstance(val, str) and val.startswith("**")
            if b:
                val = val.strip("*")
            fmt_cell(table.cell(r_idx + 1, c_idx), str(val), bold=b)
    set_table_borders(table)
    set_header_shading(table, n_cols)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    return table


# ============================================================
# PAPER SECTIONS — V3 (Complete rewrite)
# ============================================================

def write_title(doc):
    style = _try_style(doc, "Title Kinetik")
    p = doc.add_paragraph(style=style)
    p.add_run(
        "Evaluating the Reliability of LLM-Augmented Training Data "
        "for Severity-Based Hate Speech Detection in "
        "Low-Resource Javanese"
    )

def write_authors(doc):
    style_auth = _try_style(doc, "Authors Kinetik")
    p = doc.add_paragraph(style=style_auth)
    p.add_run("Mukhlis Amien")
    r = p.add_run("*1"); r.font.superscript = True
    p.add_run(", Daniel Rudiaman Sijabat")
    r = p.add_run("2"); r.font.superscript = True
    p.add_run(", Yekti Asmoro Kanthi")
    r = p.add_run("3"); r.font.superscript = True

    style_aff = _try_style(doc, "Afiliation Kinetik")
    p = doc.add_paragraph(style=style_aff)
    r = p.add_run("1,2"); r.font.superscript = True
    p.add_run("Department of Informatics, ")
    r = p.add_run("3"); r.font.superscript = True
    p.add_run("Department of Information System")
    p = doc.add_paragraph(style=style_aff)
    p.add_run("Universitas Bhinneka Nusantara, Malang, Indonesia")
    p = doc.add_paragraph(style=style_aff)
    p.add_run("*mukhlisramien@gmail.com")


def write_abstract(doc, data):
    """Write abstract in Kinetik 2-column format (Article Info | Abstract)."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    for cell in table.columns[0].cells:
        cell.width = Cm(5.0)
    for cell in table.columns[1].cells:
        cell.width = Cm(12.0)

    # Remove table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # LEFT CELL: Article Info
    left_cell = table.cell(0, 0)
    left_cell.paragraphs[0].clear()
    def _left_line(label, value=""):
        p = left_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        if value:
            r = p.add_run(f"{label}\n"); r.bold = True; r.font.size = Pt(8); r.font.name = "Arial"
            r2 = p.add_run(value); r2.font.size = Pt(8); r2.font.name = "Arial"; r2.italic = True
        else:
            r = p.add_run(label); r.bold = True; r.font.size = Pt(9); r.font.name = "Arial"
    _left_line("Article Info:")
    _left_line("Article history:", "")
    _left_line("")
    _left_line("Keywords:")
    for kw in ["Data Augmentation", "Evaluation Bias", "Low-Resource NLP",
                "Pre-trained Language Model", "Severity Classification"]:
        p = left_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(kw); r.font.size = Pt(8); r.font.name = "Arial"; r.italic = True

    # RIGHT CELL: Abstract
    right_cell = table.cell(0, 1)
    right_cell.paragraphs[0].clear()
    p = right_cell.add_paragraph()
    r = p.add_run("Abstract"); r.bold = True; r.font.size = Pt(10); r.font.name = "Arial"
    p.paragraph_format.space_after = Pt(4)

    abstract_text = (
        "The increasing adoption of large language model (LLM)-generated data for training "
        "hate speech classifiers in low-resource languages raises a critical but underexplored "
        "question: do performance metrics evaluated on mixed manual-synthetic test sets reliably "
        "reflect real-world detection capability? This study addressed this question through "
        "severity-based hate speech detection in Javanese, a low-resource language with over "
        "80 million speakers. We constructed a dataset of 9,775 annotated samples across four "
        "severity levels, comprising 46.4% manually annotated and 53.6% LLM-generated texts. "
        "Five models were evaluated using a dual-track evaluation protocol that separately "
        "measured performance on the full mixed test set and a manual-only subset. Results "
        "revealed substantial evaluation bias: XLM-RoBERTa Large achieved 80.26% F1-Macro on "
        "the full test set but only 53.89% on manual-only data, while synthetic test data "
        "yielded 99.41%. A source distinguishability analysis showed that a simple classifier "
        "could differentiate synthetic from manual texts with 97.26% F1, indicating "
        "fundamentally different distributions. Augmentation ratio experiments confirmed that "
        "manual-only performance remained flat at approximately 47% regardless of synthetic data "
        "volume, proving the inflation is a pure evaluation artifact. These findings demonstrate "
        "that LLM-augmented evaluation substantially overestimates real-world hate speech "
        "detection performance and highlight the need for separate manual-only evaluation."
    )
    p = right_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(abstract_text)
    r.font.size = Pt(9); r.font.name = "Arial"


def write_introduction(doc):
    add_heading_main(doc, "1. Introduction")

    add_body(doc,
        "Online hate speech has become a pervasive global phenomenon that threatens social "
        "cohesion, with platforms processing billions of posts daily across diverse languages "
        "and cultures [1]. Automated detection systems based on natural language processing (NLP) "
        "have emerged as essential tools for content moderation at scale [2]. However, the vast "
        "majority of research and resources have concentrated on English and other high-resource "
        "languages, leaving speakers of regional and minority languages inadequately protected [3]."
    )

    add_body(doc,
        "Indonesia presents a particularly acute case of this disparity. As the world's fourth "
        "most populous nation with over 700 living languages [4], Indonesia's linguistic diversity "
        "far exceeds the capacity of current NLP systems. While hate speech detection for formal "
        "Indonesian has received growing attention [5][6], regional languages such as Javanese "
        "\u2014 spoken by approximately 82 million people \u2014 remain severely underserved. Only "
        "limited work has addressed Javanese hate speech detection, exclusively at the binary "
        "(hate/not-hate) level [7], leaving severity-based classification entirely unexplored."
    )

    add_body(doc,
        "To overcome the scarcity of annotated data in low-resource languages, researchers have "
        "increasingly turned to large language model (LLM)-based data augmentation [8][9]. This "
        "approach uses LLMs to generate synthetic training samples that expand small seed datasets, "
        "with studies reporting improvements of 3-26% in F1 scores [8]. For hate speech detection "
        "specifically, machine-generated data has been shown to improve classifier performance "
        "[11][12][13]. However, a fundamental methodological question remains largely unaddressed: "
        "when synthetic data constitutes a significant portion of both training and test sets, "
        "do reported performance metrics accurately reflect a model's ability to detect "
        "naturally-occurring hate speech?"
    )

    # GAP ANALYSIS paragraph
    add_body(doc,
        "Despite the growing adoption of LLM augmentation in low-resource hate speech detection, "
        "no prior study has systematically evaluated whether performance metrics computed on mixed "
        "(manual + synthetic) test sets reliably indicate real-world detection capability. This gap "
        "is critical: if augmented test data inflates reported metrics, the field may overestimate "
        "progress in low-resource hate speech detection, potentially leading to premature deployment "
        "of inadequate systems. Prior work on evaluation pitfalls in NLP has demonstrated that "
        "standard evaluation practices can produce misleadingly optimistic estimates [14][15], yet "
        "this concern has not been specifically investigated in the context of LLM-augmented hate "
        "speech datasets."
    )

    add_body(doc,
        "This study addresses three research questions: (RQ1) How effective are transformer-based "
        "models compared to traditional machine learning approaches for severity-based Javanese "
        "hate speech classification? (RQ2) Does LLM data augmentation inflate reported performance "
        "metrics when synthetic samples appear in the test set? (RQ3) What characteristics of "
        "synthetic data enable models to achieve near-perfect performance on synthetic test samples "
        "while struggling with manual ones?"
    )

    add_body(doc,
        "The contributions of this work are threefold. First, we present the first severity-level "
        "(four-class) hate speech classification dataset for Javanese, comprising 9,775 annotated "
        "samples. Second, we propose and apply a dual-track evaluation protocol that separately "
        "measures performance on manual-only and full (mixed) test sets, revealing a 45.52 "
        "percentage-point F1 gap attributable to augmentation bias. Third, we provide empirical "
        "evidence that LLM-generated data is distributionally distinct from manual annotations "
        "(97.26% source distinguishability), explaining why augmented evaluation is unreliable."
    )

    # Literature review integrated into Introduction (Kinetik template requirement)
    add_body(doc,
        "Hate speech detection has been extensively studied for high-resource languages, with "
        "Davidson et al. [16] establishing a benchmark of 24,802 English tweets. For Indonesian, "
        "Alfina et al. [5] created a dataset of 713 tweets for binary classification, Ibrohim "
        "and Budi [6] expanded to 13,169 multi-label tweets, and Fauzi and Yuniarti [10] explored "
        "ensemble methods. Susanto et al. [17] introduced IndoToxic2024 with 43,692 entries. For "
        "Javanese, Putri et al. [7] conducted preliminary binary classification, and Pamungkas "
        "and Chiril [18] addressed code-mixed Indonesian-Javanese hate speech. Table 1 "
        "contextualizes our work within this landscape."
    )

    # Table 1: Comparison with Prior Work
    create_table(doc,
        headers=["Study", "Language", "Classes", "Dataset", "Method", "F1 (%)", "Metric"],
        rows=[
            ["Alfina [5]", "Indonesian", "2", "713", "RFDT", "93.5", "F-measure"],
            ["Davidson [16]", "English", "3", "24,802", "LR+TF-IDF", "90.0*", "weighted"],
            ["Ibrohim [6]", "Indonesian", "multi", "13,169", "RFDT+LP", "77.0**", "accuracy"],
            ["Putri [7]", "Javanese", "2", "~2,500", "SVM", ">60", "F-measure"],
            ["Susanto [17]", "Indonesian", "7", "43,692", "IndoBERTweet", "78.0", "macro-F1"],
            ["**Ours (full)**", "**Javanese**", "**4**", "**9,775**", "**XLM-R**", "**80.26**", "**macro-F1**"],
            ["**Ours (manual)**", "**Javanese**", "**4**", "**4,538**", "**XLM-R**", "**53.89**", "**macro-F1**"],
        ],
        caption_text=(
            "Table 1. Comparison with Prior Work on Hate Speech Detection. "
            "*F1-weighted; **accuracy metric; not directly comparable to macro-F1."
        ),
    )

    add_body(doc,
        "Data augmentation using language models has gained traction for hate speech detection. "
        "Hartvigsen et al. [11] demonstrated with ToxiGen that machine-generated data could "
        "improve classifier training. Vidgen et al. [12] proposed dynamic dataset generation, "
        "and Juuti et al. [13] found GPT-2 augmented data enabled shallow classifiers to approach "
        "BERT-level performance. Jahan et al. [19] found GPT-3 augmentation improved F1 by 1.4%. "
        "However, none of these studies evaluated whether synthetic data in the test set inflates "
        "metrics. Gorman and Bedrick [14] and Sogaard et al. [15] showed that standard evaluation "
        "splits produce misleadingly optimistic estimates, recommending multiple independent test "
        "sets. Cross-lingual models such as XLM-RoBERTa [20] and Indonesian-specific IndoBERT [21] "
        "have shown promise for low-resource settings, though Aji et al. [4] noted that even large "
        "models may underperform on truly low-resource varieties like Javanese."
    )


def write_research_method(doc, data):
    add_heading_main(doc, "2. Research Method")

    # 2.1. Dataset Construction
    add_subheading(doc, "2.1. Dataset Construction")
    add_body(doc,
        "The dataset was constructed in two phases: manual annotation and LLM augmentation."
    )
    add_body(doc,
        "Manual annotation. A total of 4,538 Javanese texts were collected from Twitter and "
        "Instagram, then annotated by three native Javanese speakers into four severity levels: "
        "(0) Not Hate Speech, (1) Light \u2014 subtle stereotyping or microaggressions, "
        "(2) Moderate \u2014 explicit prejudice without dehumanization, and (3) Severe \u2014 "
        "dehumanizing rhetoric, calls for violence, or slurs. Inter-annotator agreement was "
        "moderate-to-substantial (Cohen's kappa = 0.72), with majority voting resolving "
        "disagreements. Table 2 presents severity level definitions with examples."
    )

    create_table(doc,
        headers=["Level", "Label", "Example"],
        rows=[
            ["0", "Not Hate", "Aku mikir umume wong setuju yen kabeh manungsa padha."],
            ["1", "Light", "Wah, ibu iki koyo detektif, sak klebatan motor tamu wae wis kudu takon."],
            ["2", "Moderate", "Wong asing teka mung gawe masalah ing kampung kita."],
            ["3", "Severe", "Kabeh keturunan iku mata duitan, kudu dipeksa lunga."],
        ],
        caption_text="Table 2. Severity Level Definitions with Examples. Examples are representative and may be abbreviated.",
    )

    add_body(doc,
        "LLM augmentation. To address class imbalance and data scarcity, 5,237 additional "
        "samples were generated using DeepSeek-Coder-V2 (236B parameters) through prompted "
        "generation with severity-specific instructions. Generated texts underwent quality "
        "filtering and human verification on a subset of 500 samples (kappa = 0.72). After "
        "cleaning (removing duplicates and short texts), the final dataset comprised 9,775 "
        "samples: 4,538 manual (46.4%) and 5,237 synthetic (53.6%). Each sample was tagged "
        "with its source to enable source-aware evaluation."
    )

    # Table 3: Dataset Composition
    dq = data.get("data_quality")
    pcd = dq["results"]["per_class_distribution"] if dq else {}
    rows_t3 = []
    labels = ["Not Hate", "Light", "Moderate", "Severe"]
    for i, label in enumerate(labels):
        cls = pcd.get(f"class_{i}", {})
        m = cls.get("manual", "?")
        s = cls.get("synthetic", "?")
        t = cls.get("total", "?")
        sr = f"{cls.get('synthetic_ratio', 0):.0%}" if cls else "?"
        rows_t3.append([f"{i} - {label}", str(m), str(s), str(t), sr])
    rows_t3.append(["Total", "4,538", "5,237", "9,775", "54%"])

    create_table(doc,
        headers=["Class", "Manual", "Synthetic", "Total", "Synth %"],
        rows=rows_t3,
        caption_text="Table 3. Dataset Composition by Class and Source.",
    )

    # 3.2 Dual-Track Evaluation Protocol
    add_subheading(doc, "2.2. Dual-Track Evaluation Protocol")
    add_body(doc,
        "To assess whether LLM augmentation inflates evaluation metrics, we employed a "
        "dual-track evaluation protocol. The dataset was split into training (80%), validation "
        "(10%), and test (10%) sets using stratified sampling (seed=42). The source "
        "(manual/synthetic) of each sample was tracked through the split."
    )
    add_body(doc,
        "Evaluation was conducted on: (1) the full test set (978 samples, approximately 46% "
        "manual and 54% synthetic), representing standard mixed evaluation; and (2) the "
        "manual-only test subset (451 samples), reflecting real-world performance on "
        "naturally-occurring texts. Additionally, an augmentation ratio study was conducted, "
        "training models with 0%, 25%, 50%, 75%, and 100% of available synthetic data while "
        "always evaluating on the manual-only test subset."
    )

    # 3.3 Models
    add_subheading(doc, "2.3. Models")
    add_body(doc,
        "Five models were evaluated, spanning traditional machine learning and transformer "
        "architectures. Baselines included SVM with TF-IDF features [22] (linear kernel, "
        "max 10,000 features, unigrams and bigrams) and Logistic Regression with identical "
        "feature extraction. Transformer models included IndoBERT [21] "
        "(indobenchmark/indobert-base-p1, 110M parameters), IndoBERT with label smoothing "
        "(epsilon=0.1) [23][24], and XLM-RoBERTa Large [20] (xlm-roberta-large, 559M parameters)."
    )

    # 3.4 Experimental Setup
    add_subheading(doc, "2.4. Experimental Setup")
    add_body(doc,
        "All transformer models were fine-tuned for 5 epochs with learning rate 2e-5, weight "
        "decay 0.01, warmup ratio 0.1, and batch sizes of 16 (IndoBERT) or 8 (XLM-R Large). "
        "Maximum sequence length was 128 tokens. Training used mixed precision (FP16) on an "
        "NVIDIA RTX 4080 GPU. The best checkpoint was selected based on validation F1-Macro. "
        "Five-fold stratified cross-validation was conducted for baseline models to verify "
        "robustness."
    )


def write_results_discussion(doc, data):
    add_heading_main(doc, "3. Results and Discussion")

    comp = data.get("comparative")
    manual_only = data.get("manual_only")
    ratio = data.get("ratio")
    dq = data.get("data_quality")
    error = data.get("error_analysis")
    two_stage = data.get("two_stage")
    cv = data.get("cv_baselines")

    # 4.1 Overall Performance
    add_subheading(doc, "3.1. Overall Performance Comparison")
    add_body(doc,
        "Table 4 presents the performance of all models on both the full test set and the "
        "manual-only test subset."
    )

    # Build Table 4 from data
    models = [
        ("LR + TF-IDF", "lr_tfidf"),
        ("SVM + TF-IDF", "svm_tfidf"),
        ("IndoBERT", "indobert"),
        ("IndoBERT + LS", "indobert_ls"),
        ("XLM-R Large", "xlmr_large"),
    ]
    rows_t4 = []
    for name, key in models:
        full_f1 = "?"
        man_f1 = "?"
        if manual_only and key in manual_only.get("models", {}):
            m = manual_only["models"][key]
            full_f1 = f"{m['full_test']['f1_macro']:.2f}"
            man_f1 = f"{m['manual_only']['f1_macro']:.2f}"
        gap = "?"
        try:
            gap = f"{float(full_f1) - float(man_f1):.2f}"
        except (ValueError, TypeError):
            pass
        bold = "**" if key == "xlmr_large" else ""
        rows_t4.append([f"{bold}{name}{bold}", f"{bold}{full_f1}{bold}",
                        f"{bold}{man_f1}{bold}", f"{bold}{gap}{bold}"])

    create_table(doc,
        headers=["Model", "Full Test (978)", "Manual-Only (451)", "Gap (pp)"],
        rows=rows_t4,
        caption_text="Table 4. Model Performance on Full and Manual-Only Test Sets (F1-Macro %).",
    )

    add_body(doc,
        "All models showed a consistent gap exceeding 26 percentage points between full and "
        "manual-only test performance. XLM-RoBERTa Large achieved the best performance on both "
        "evaluations (80.26% full, 53.89% manual), followed by IndoBERT with label smoothing "
        "(77.36% full, 49.30% manual). Notably, the traditional SVM baseline (48.55% manual) "
        "performed comparably to IndoBERT (45.27%), suggesting that transformer scale alone does "
        "not guarantee superior real-world hate speech detection for this task."
    )

    # CV results + statistical significance
    if cv:
        svm_full = cv.get("svm_full", {})
        svm_man = cv.get("svm_manual_only", {})
        add_body(doc,
            f"Five-fold cross-validation confirmed consistency: SVM achieved "
            f"{svm_full.get('mean', '?')}% +/- {svm_full.get('std', '?')}% on the full dataset "
            f"and {svm_man.get('mean', '?')}% +/- {svm_man.get('std', '?')}% on manual-only data. "
            f"The approximately 30-point gap held across all folds and was statistically "
            f"significant (independent t-test, p < 0.001, Cohen's d > 33), confirming "
            f"that the evaluation inflation is not an artifact of a single data split."
        )

    # 4.2 Comparison with Prior Work
    add_subheading(doc, "3.2. Comparison with Prior Work")
    add_body(doc,
        "Direct comparison with prior work is challenging due to differences in language, number "
        "of classes, and dataset composition (Table 1). Our full-test F1 of 80.26% appears "
        "competitive with or superior to prior Indonesian hate speech detection work. However, "
        "this comparison is misleading: our manual-only F1 of 53.89% reveals that the apparent "
        "competitiveness is an artifact of synthetic data in the test set. This finding has "
        "broad implications: prior studies using LLM-augmented data [11][12][13] that report only "
        "mixed-set metrics may similarly overestimate real-world performance."
    )

    # 4.3 Augmentation Bias Analysis
    add_subheading(doc, "3.3. Augmentation Bias Analysis")
    add_body(doc,
        "The central finding of this study is a substantial performance gap between evaluation "
        "on synthetic and manual test data. XLM-RoBERTa Large achieved 99.41% F1 on the "
        "synthetic-only test subset (527 samples) versus 53.89% on the manual-only subset "
        "(451 samples) \u2014 a gap of 45.52 percentage points."
    )

    # Table 5: Ratio Study
    if ratio:
        rows_t5 = []
        for r_pct in [0, 25, 50, 75, 100]:
            rk = f"ratio_{r_pct}"
            rd = ratio["results"].get(rk, {})
            svm_m = rd.get("svm", {}).get("manual_test", {}).get("f1_macro", "?")
            svm_f = rd.get("svm", {}).get("full_test", {}).get("f1_macro", "?")
            ib_m = rd.get("indobert", {}).get("manual_test", {}).get("f1_macro", "?")
            ib_f = rd.get("indobert", {}).get("full_test", {}).get("f1_macro", "?")
            rows_t5.append([f"{r_pct}%", str(svm_m), str(svm_f), str(ib_m), str(ib_f)])

        create_table(doc,
            headers=["Synth Ratio", "SVM Manual", "SVM Full", "IndoBERT Manual", "IndoBERT Full"],
            rows=rows_t5,
            caption_text="Table 5. Augmentation Ratio Study - F1-Macro (%) on Manual-Only and Full Test Sets.",
        )

    add_figure(doc, "figure_augmentation_ratio.png",
        "Figure 1. Impact of synthetic data ratio on F1-Macro for manual-only (a) and full (b) test sets."
    )

    add_body(doc,
        "Manual-only F1 remained essentially flat (approximately 47%) regardless of augmentation "
        "ratio for both SVM and IndoBERT (Table 5, Figure 1a). In contrast, full-test F1 jumped "
        "approximately 35 percentage points with just 25% synthetic data inclusion (Figure 1b). "
        "This demonstrated that the performance improvement from augmentation was predominantly "
        "an evaluation artifact: synthetic data in the test set inflates metrics without "
        "meaningfully improving detection of naturally-occurring hate speech."
    )

    add_body(doc,
        "An additional observation was that IndoBERT trained exclusively on manual data (0% "
        "synthetic, F1=48.75%) did not collapse, unlike XLM-RoBERTa Large which degraded to "
        "14.64% under the same condition. This suggests that smaller models are more robust to "
        "limited training data, while larger models require augmentation simply to avoid training "
        "collapse rather than to genuinely improve hate speech understanding."
    )

    # 4.4 Data Quality Analysis
    add_subheading(doc, "3.4. Data Quality Analysis")
    if dq:
        dist_f1 = dq["results"]["source_distinguishability"]["cv_f1_macro_mean"]
        add_body(doc,
            f"To understand why models achieved near-perfect performance on synthetic data, we "
            f"conducted a distributional analysis. A logistic regression classifier trained on "
            f"TF-IDF features to distinguish synthetic from manual texts achieved "
            f"{dist_f1}% +/- {dq['results']['source_distinguishability']['cv_f1_macro_std']}% "
            f"F1 in 5-fold cross-validation. This indicates fundamentally different distributions, "
            f"enabling models to learn source-specific shortcuts rather than genuine hate speech "
            f"features."
        )

    # Table 6: Data quality
    if dq:
        ms = dq["results"]["manual_stats"]
        ss = dq["results"]["synthetic_stats"]
        sp = dq["results"]["sentence_patterns"]
        vo = dq["results"]["vocabulary_overlap"]
        create_table(doc,
            headers=["Feature", "Manual", "Synthetic"],
            rows=[
                ["Mean word count", f"{ms['word_count_mean']:.1f}", f"{ss['word_count_mean']:.1f}"],
                ["Vocabulary size", str(ms["vocab_size"]), str(ss["vocab_size"])],
                ["Type-token ratio", f"{ms['type_token_ratio']:.3f}", f"{ss['type_token_ratio']:.3f}"],
                ["Jaccard vocab overlap", "--", f"{vo['jaccard_similarity']:.1%}"],
                ["Starts with capital", f"{sp['manual']['starts_with_capital']:.1%}", f"{sp['synthetic']['starts_with_capital']:.1%}"],
                ["Ends with period", f"{sp['manual']['ends_with_period']:.1%}", f"{sp['synthetic']['ends_with_period']:.1%}"],
                ["All lowercase texts", f"{sp['manual']['all_lowercase']:.1%}", f"{sp['synthetic']['all_lowercase']:.1%}"],
            ],
            caption_text="Table 6. Distributional Differences Between Manual and Synthetic Data.",
        )

    add_figure(doc, "figure_data_quality.png",
        "Figure 2. Surface pattern differences between manual and synthetic data.",
        width=Inches(4.5),
    )

    add_body(doc,
        "The synthetic data exhibited highly uniform surface structure: all texts started with "
        "a capital letter, 95.8% ended with periods, and none appeared in all lowercase \u2014 "
        "patterns absent in natural social media text. The vocabulary overlap between sources was "
        "only 14.6% (Jaccard similarity). The most distinctive synthetic indicators were formal "
        "political terms (rakyat/people, pemerintah/government), while manual indicators were "
        "informal personal terms (wadon/woman, aku/I, ireng/black), suggesting the LLM generated "
        "texts in a formal register inconsistent with real social media hate speech."
    )

    # 4.5 Error Analysis
    add_subheading(doc, "3.5. Error Analysis")
    if error:
        xlmr = error["results"].get("xlmr_large", {})
        pc = xlmr.get("per_class_errors", {})
        add_body(doc,
            f"Error analysis of XLM-RoBERTa Large predictions on the manual-only test set "
            f"({xlmr.get('total_errors', '?')} errors out of 451 samples, error rate "
            f"{xlmr.get('error_rate', '?')}%) revealed that {xlmr.get('adjacent_error_ratio', '?')}% "
            f"of misclassifications involved adjacent severity classes. The most common "
            f"error pairs were Not Hate to Light ({xlmr['top_error_pairs'][0]['pct_of_errors']}% "
            f"of errors) and Light to Not Hate ({xlmr['top_error_pairs'][1]['pct_of_errors']}%), "
            f"confirming that the boundary between neutral speech and light hate speech is the "
            f"most subjective and challenging."
        )

    if error:
        xlmr = error["results"].get("xlmr_large", {})
        pc = xlmr.get("per_class_errors", {})
        add_body(doc,
            f"Per-class accuracy varied substantially on manual data: Not Hate achieved "
            f"{pc.get('Not Hate', {}).get('accuracy', '?')}% (most distinct class), "
            f"Light {pc.get('Light', {}).get('accuracy', '?')}%, "
            f"Severe {pc.get('Severe', {}).get('accuracy', '?')}%, and "
            f"Moderate only {pc.get('Moderate', {}).get('accuracy', '?')}%. "
            f"The low accuracy for Moderate aligns with its having the highest synthetic "
            f"proportion (75% in Table 3), suggesting that the model's understanding of "
            f"moderate-severity hate speech was dominated by synthetic patterns rather than "
            f"genuine linguistic markers. The Moderate class was also the smallest in the manual "
            f"test set (61 samples), compounding the difficulty."
        )

    add_figure(doc, "figure3_confusion_matrix_manual.png",
        "Figure 3. Confusion matrix of XLM-RoBERTa Large on manual-only test set (451 samples).",
        width=Inches(4.0),
    )

    if error:
        xlmr = error["results"].get("xlmr_large", {})
        conf = xlmr.get("confidence_analysis", {})
        agreement = error["results"].get("cross_model_agreement", {})
        add_body(doc,
            f"Confidence analysis revealed overconfidence on incorrect predictions: mean "
            f"confidence was {conf.get('correct_mean_confidence', '?'):.3f} for correct "
            f"predictions but still {conf.get('error_mean_confidence', '?'):.3f} for errors, "
            f"indicating that prediction confidence is not a reliable indicator of correctness. "
            f"Cross-model agreement analysis between SVM and XLM-RoBERTa Large showed "
            f"complementary error patterns: SVM correctly classified 52 samples that XLM-R missed, "
            f"and vice versa for 74 samples, yielding a theoretical ensemble upper bound of "
            f"{agreement.get('ensemble_potential_pct', '?')}%, substantially higher than either "
            f"model alone."
        )

    # 4.6 Two-Stage Training
    if two_stage:
        add_subheading(doc, "3.6. Two-Stage Training")

        # Build Table 7
        rows_ts = []
        for mk, name in [("indobert", "IndoBERT"), ("xlmr_large", "XLM-R Large")]:
            if mk in two_stage.get("results", {}):
                r = two_stage["results"][mk]
                s1m = r["stage1_synthetic_only"]["manual_test"]["f1_macro"]
                s2m = r["stage2_manual_finetune"]["manual_test"]["f1_macro"]
                s2f = r["stage2_manual_finetune"]["full_test"]["f1_macro"]
                rows_ts.append([name, f"{s1m:.2f}", f"{s2m:.2f}", f"{s2f:.2f}"])

        create_table(doc,
            headers=["Model", "Stage 1 Manual", "Stage 2 Manual", "Stage 2 Full"],
            rows=rows_ts,
            caption_text="Table 7. Two-Stage Training Results (F1-Macro %). Stage 1: synthetic-only pre-training. Stage 2: manual fine-tuning.",
        )

        add_body(doc,
            "To investigate whether curriculum learning could mitigate augmentation bias, we "
            "employed a two-stage training approach: Stage 1 fine-tuned models on synthetic data "
            "only (3 epochs, lr=2e-5), followed by Stage 2 fine-tuning on manual data (5 epochs, "
            "lr=5e-6). As shown in Table 7, two-stage training did not improve manual-only F1 "
            "beyond single-stage training for either model. XLM-RoBERTa Large achieved 52.16% "
            "(versus 53.89% single-stage, a decrease of 1.73 points), while IndoBERT achieved "
            "45.34% (versus 45.27%, essentially unchanged)."
        )

        add_body(doc,
            "However, two-stage training did prevent XLM-RoBERTa Large from training collapse: "
            "when trained directly on manual-only data without any synthetic samples, XLM-R "
            "collapsed to 14.64% F1, whereas the two-stage approach achieved 52.16% by first "
            "learning general patterns from synthetic data. This suggests that curriculum learning "
            "serves a similar function to standard augmentation \u2014 preventing collapse of "
            "overparameterized models \u2014 but does not improve the model's ability to detect "
            "hate speech in natural text beyond what single-stage training achieved."
        )

    # 4.7 Limitations
    add_subheading(doc, "3.7. Limitations")
    add_body(doc,
        "Several limitations should be acknowledged. First, the inter-annotator agreement "
        "(kappa=0.72, moderate-to-substantial) indicates inherent label subjectivity, "
        "particularly for boundary cases between adjacent severity levels [25][26]. Second, "
        "the dataset was collected from Twitter and Instagram only, potentially missing hate "
        "speech patterns from other platforms. Third, this study focused on text-only analysis, "
        "while hate speech increasingly involves multimodal content. Finally, while our findings "
        "demonstrate evaluation bias from LLM augmentation, the optimal mitigation strategy for "
        "low-resource hate speech detection remains an open question."
    )


def write_conclusion(doc, data):
    add_heading_main(doc, "4. Conclusion")

    add_body(doc,
        "This study investigated the reliability of LLM-augmented evaluation for severity-based "
        "hate speech detection in Javanese through three research questions. Regarding RQ1, "
        "XLM-RoBERTa Large achieved the highest F1-Macro (53.89% on manual data), but "
        "traditional SVM performed comparably (48.55%), suggesting that transformer scale alone "
        "does not guarantee superior real-world performance for this low-resource task. Regarding "
        "RQ2, LLM augmentation substantially inflated reported metrics: the same model scored "
        "80.26% on the mixed test set but only 53.89% on manual-only data, a 26.37-point gap. "
        "The augmentation ratio study confirmed this inflation was consistent across models, "
        "with manual-only F1 remaining flat at approximately 47% regardless of synthetic data "
        "volume. Regarding RQ3, source distinguishability analysis revealed that a simple "
        "classifier could differentiate synthetic from manual texts with 97.26% F1. The synthetic "
        "data differed fundamentally in vocabulary (14.6% Jaccard overlap), sentence length "
        "(10.8 vs 16.5 words), and surface patterns (100% vs 74.6% capitalization), explaining "
        "why models learned source-specific shortcuts rather than genuine hate speech features."
    )

    add_body(doc,
        "These findings carry three practical implications for the low-resource NLP community. "
        "First, studies employing LLM data augmentation should adopt dual-track evaluation, "
        "reporting both mixed-set and manual-only metrics to provide realistic performance "
        "estimates. Second, augmentation primarily prevents training collapse of large models "
        "rather than genuinely improving hate speech understanding, as confirmed by both the "
        "ratio study and the negative two-stage training results. Third, the 97.26% source "
        "distinguishability indicates that improving synthetic data quality to better match "
        "natural text distributions is essential before augmentation can yield genuine "
        "performance gains. Future work should explore ensemble methods (theoretical upper "
        "bound: 69.18%), improved generation prompts that produce more naturalistic text, "
        "and larger manually-annotated Javanese hate speech datasets."
    )


def write_acknowledgement(doc):
    add_heading_main(doc, "Acknowledgement")
    add_body(doc,
        "The authors thank the anonymous reviewers for their constructive feedback and the "
        "Javanese-speaking annotators who contributed to dataset construction."
    )


def write_references(doc):
    add_heading_main(doc, "References")
    refs = [
        '[1] R. Ramos et al., "A comprehensive review on hate speech detection using NLP," IEEE Access, 2024.',
        '[2] M. A. Hedderich et al., "A survey on recent approaches for natural language processing in low-resource scenarios," in Proc. NAACL, 2021, pp. 2545-2568.',
        '[3] S. Cahyawijaya et al., "NusaCrowd: Open source initiative for Indonesian NLP resources," in Findings of ACL, 2023.',
        '[4] A. F. Aji et al., "One country, 700+ languages: NLP challenges for underrepresented languages and dialects in Indonesia," in Proc. 60th ACL, 2022, pp. 7226-7249.',
        '[5] I. Alfina et al., "Hate speech detection in the Indonesian language: A dataset and preliminary study," in Proc. ICACSIS, 2017, pp. 233-238.',
        '[6] M. O. Ibrohim and I. Budi, "Multi-label hate speech and abusive language detection in Indonesian Twitter," in Proc. ALW3, ACL, 2019, pp. 46-57.',
        '[7] S. D. A. Putri et al., "Abusive language and hate speech detection for Javanese and Sundanese languages," in Proc. WCSE, 2021, pp. 461-465.',
        '[8] B. Ding et al., "Data augmentation using LLMs: Data perspectives, learning paradigms and challenges," arXiv:2403.02990, 2024.',
        '[9] H. Rizwan et al., "Exploring conditional language model based data augmentation approaches for hate speech classification," in LNCS, vol. 12882, Springer, 2022.',
        '[10] M. A. Fauzi and A. P. Yuniarti, "Ensemble method for Indonesian tweet hate speech detection," Indonesian J. Elec. Eng. Comput. Sci., vol. 11, no. 1, pp. 294-299, 2018.',
        '[11] T. Hartvigsen et al., "ToxiGen: A large-scale machine-generated dataset for adversarial and implicit hate speech detection," in Proc. 60th ACL, 2022, pp. 3309-3326.',
        '[12] B. Vidgen et al., "Learning from the worst: Dynamically generated datasets to improve online hate detection," in Proc. 59th ACL, 2021, pp. 1667-1682.',
        '[13] M. Juuti et al., "A little goes a long way: Improving toxic language classification despite data scarcity," in Findings of EMNLP, 2020, pp. 2991-3009.',
        '[14] K. Gorman and S. Bedrick, "We need to talk about standard splits," in Proc. 57th ACL, 2019, pp. 2786-2791.',
        '[15] A. Sogaard et al., "We need to talk about random splits," in Proc. EACL, 2021, pp. 1823-1832.',
        '[16] T. Davidson et al., "Automated hate speech detection and the problem of offensive language," in Proc. ICWSM, vol. 11, 2017, pp. 512-515.',
        '[17] L. Susanto et al., "IndoToxic2024: A demographically-enriched dataset of hate speech and toxicity types for Indonesian language," arXiv:2406.19349, 2024.',
        '[18] E. W. Pamungkas and P. Chiril, "Ngalawan Ujaran Sengit: Hate speech detection in Indonesian code-mixed social media data," Lang. Resour. Eval., vol. 59, pp. 2387-2414, 2025.',
        '[19] M. S. Jahan et al., "A comprehensive study on NLP data augmentation for hate speech detection," in Findings of ACL, 2024.',
        '[20] A. Conneau et al., "Unsupervised cross-lingual representation learning at scale," in Proc. 58th ACL, 2020, pp. 8440-8451.',
        '[21] B. Wilie et al., "IndoNLU: Benchmark and resources for evaluating Indonesian natural language understanding," in Proc. AACL-IJCNLP, 2020.',
        '[22] T. Joachims, "Text categorization with support vector machines: Learning with many relevant features," in Proc. ECML, 1998, pp. 137-142.',
        '[23] C. Szegedy et al., "Rethinking the inception architecture for computer vision," in Proc. CVPR, 2016, pp. 2818-2826.',
        '[24] R. Muller et al., "When does label smoothing help?," in Proc. NeurIPS, 2019, pp. 4694-4703.',
        '[25] L. Aroyo and C. Welty, "Truth is a lie: Crowd truth and the seven myths of human annotation," AI Mag., vol. 36, no. 1, pp. 15-24, 2015.',
        '[26] S. Paun et al., "Comparing Bayesian models of annotation," Trans. ACL, vol. 6, pp. 571-585, 2018.',
        '[27] J. Devlin et al., "BERT: Pre-training of deep bidirectional transformers for language understanding," in Proc. NAACL-HLT, 2019, pp. 4171-4186.',
        '[28] T. G. Dietterich, "Ensemble methods in machine learning," in MCS, LNCS, vol. 1857, 2000, pp. 1-15.',
    ]
    ref_style = _try_style(doc, "References kinetik")
    for ref in refs:
        p = doc.add_paragraph(style=ref_style)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(ref)
        run.font.size = Pt(8)
        run.font.name = "Arial"


# ============================================================
# MAIN
# ============================================================
def main():
    print("=" * 60)
    print("GENERATING KINETIK PAPER V3")
    print("=" * 60)

    # Load all results
    print("\nLoading results...")
    data = {
        "comparative": load_json("comparative_results"),
        "manual_only": load_json("manual_only_results"),
        "augmentation": load_json("augmentation_impact"),
        "ratio": load_json("augmentation_ratio_results"),
        "data_quality": load_json("data_quality_analysis"),
        "error_analysis": load_json("error_analysis"),
        "two_stage": load_json("two_stage_results"),
        "cv_baselines": None,
        "ablation": load_json("ablation_results"),
        "multiseed": load_json("multi_seed_results"),
        "cleaning": load_cleaning(),
    }

    # Load CV baselines
    cv_path = RESULTS / "cv_baselines.json"
    if cv_path.exists():
        with open(cv_path, encoding="utf-8") as f:
            data["cv_baselines"] = json.load(f)

    # Generate confusion matrix figure
    print("\nGenerating figures...")
    try:
        generate_confusion_matrix(data["manual_only"])
    except Exception as e:
        print(f"  [WARN] Could not generate confusion matrix: {e}")

    # Open template
    print(f"\nTemplate: {TEMPLATE}")
    doc = Document(str(TEMPLATE))
    clear_body(doc)
    update_footers(doc)

    # Write paper
    print("\nWriting sections...")
    write_title(doc)
    print("  Title")
    write_authors(doc)
    print("  Authors")
    write_abstract(doc, data)
    print("  Abstract")
    write_introduction(doc)
    print("  Introduction (with literature review)")
    write_research_method(doc, data)
    print("  Research Method")
    write_results_discussion(doc, data)
    print("  Results and Discussion")
    write_conclusion(doc, data)
    print("  Conclusion")
    write_acknowledgement(doc)
    print("  Acknowledgement")
    write_references(doc)
    print("  References")

    # Save
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"\nPaper saved: {OUTPUT}")
    print("Done!")


def generate_confusion_matrix(manual_only):
    """Generate confusion matrix figure."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import seaborn as sns
    except ImportError:
        return

    if not manual_only:
        return

    cm_data = manual_only["models"]["xlmr_large"]["manual_only"]["confusion_matrix"]
    cm = np.array(cm_data)
    labels = ["Not Hate", "Light", "Moderate", "Severe"]

    fig, ax = plt.subplots(figsize=(6, 5))
    sns.heatmap(cm, annot=True, fmt="d", cmap="Blues",
                xticklabels=labels, yticklabels=labels, ax=ax)
    ax.set_xlabel("Predicted", fontsize=11)
    ax.set_ylabel("Actual", fontsize=11)
    ax.set_title("XLM-RoBERTa Large - Manual-Only Test (451 samples)", fontsize=12)
    plt.tight_layout()
    out = FIGURES / "figure3_confusion_matrix_manual.png"
    out.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(str(out), dpi=150)
    plt.close(fig)
    print(f"  Generated: {out}")


if __name__ == "__main__":
    main()
