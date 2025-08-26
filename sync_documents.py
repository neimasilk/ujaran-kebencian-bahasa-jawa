#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script untuk sinkronisasi draft paper dengan template jurnal
Memberikan panduan praktis dan checklist untuk penyesuaian format

Author: Assistant
Date: 2024
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.table import Table

def analyze_document_structure(doc_path, doc_name):
    """
    Menganalisis struktur dokumen
    
    Args:
        doc_path (str): Path ke dokumen
        doc_name (str): Nama dokumen untuk display
    
    Returns:
        dict: Informasi struktur dokumen
    """
    try:
        doc = Document(doc_path)
        
        # Hitung statistik dasar
        total_paragraphs = len(doc.paragraphs)
        total_tables = len(doc.tables)
        
        # Analisis styles yang digunakan
        used_styles = set()
        content_sections = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                used_styles.add(para.style.name)
                
                # Identifikasi bagian-bagian penting
                text_lower = text.lower()
                if any(keyword in text_lower for keyword in ['abstract', 'intisari']):
                    content_sections.append(('ABSTRACT', text[:100], para.style.name))
                elif any(keyword in text_lower for keyword in ['keywords', 'kata kunci']):
                    content_sections.append(('KEYWORDS', text[:100], para.style.name))
                elif any(keyword in text_lower for keyword in ['introduction', 'pendahuluan']):
                    content_sections.append(('INTRODUCTION', text[:100], para.style.name))
                elif any(keyword in text_lower for keyword in ['method', 'metodologi']):
                    content_sections.append(('METHOD', text[:100], para.style.name))
                elif any(keyword in text_lower for keyword in ['result', 'hasil']):
                    content_sections.append(('RESULTS', text[:100], para.style.name))
                elif any(keyword in text_lower for keyword in ['conclusion', 'kesimpulan']):
                    content_sections.append(('CONCLUSION', text[:100], para.style.name))
                elif any(keyword in text_lower for keyword in ['reference', 'daftar pustaka']):
                    content_sections.append(('REFERENCES', text[:100], para.style.name))
        
        return {
            'name': doc_name,
            'paragraphs': total_paragraphs,
            'tables': total_tables,
            'styles': sorted(list(used_styles)),
            'sections': content_sections
        }
        
    except Exception as e:
        print(f"Error menganalisis {doc_name}: {str(e)}")
        return None

def get_template_specifications(template_path):
    """
    Mendapatkan spesifikasi format dari template
    
    Args:
        template_path (str): Path ke template
    
    Returns:
        dict: Spesifikasi format template
    """
    try:
        doc = Document(template_path)
        
        # Page setup
        section = doc.sections[0]
        page_setup = {
            'top_margin': round(section.top_margin.inches, 2),
            'bottom_margin': round(section.bottom_margin.inches, 2),
            'left_margin': round(section.left_margin.inches, 2),
            'right_margin': round(section.right_margin.inches, 2),
            'page_width': round(section.page_width.inches, 2),
            'page_height': round(section.page_height.inches, 2)
        }
        
        # Style specifications
        style_specs = {}
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                spec = {'name': style.name}
                
                if hasattr(style, 'font') and style.font:
                    if style.font.name:
                        spec['font_name'] = style.font.name
                    if style.font.size:
                        spec['font_size'] = style.font.size.pt
                    if style.font.bold is not None:
                        spec['bold'] = style.font.bold
                    if style.font.italic is not None:
                        spec['italic'] = style.font.italic
                
                if hasattr(style, 'paragraph_format') and style.paragraph_format:
                    pf = style.paragraph_format
                    if pf.alignment is not None:
                        alignment_map = {0: 'Left', 1: 'Center', 2: 'Right', 3: 'Justify'}
                        spec['alignment'] = alignment_map.get(pf.alignment, 'Unknown')
                    if pf.space_before:
                        spec['space_before'] = pf.space_before.pt
                    if pf.space_after:
                        spec['space_after'] = pf.space_after.pt
                    if pf.line_spacing:
                        spec['line_spacing'] = pf.line_spacing
                
                style_specs[style.name] = spec
        
        return {
            'page_setup': page_setup,
            'styles': style_specs
        }
        
    except Exception as e:
        print(f"Error mendapatkan spesifikasi template: {str(e)}")
        return None

def generate_sync_guide(draft_info, template_info, template_specs):
    """
    Membuat panduan sinkronisasi yang detail
    
    Args:
        draft_info (dict): Informasi draft
        template_info (dict): Informasi template
        template_specs (dict): Spesifikasi template
    """
    print(f"\n{'='*80}")
    print("PANDUAN SINKRONISASI DRAFT PAPER DENGAN TEMPLATE JURNAL")
    print(f"{'='*80}")
    
    print(f"\nDRAFT: {draft_info['name']}")
    print(f"TEMPLATE: {template_info['name']}")
    
    # Perbandingan struktur
    print(f"\n{'='*60}")
    print("1. PERBANDINGAN STRUKTUR DOKUMEN")
    print(f"{'='*60}")
    
    print(f"\nStatistik Dokumen:")
    print(f"  Draft    : {draft_info['paragraphs']} paragraf, {draft_info['tables']} tabel")
    print(f"  Template : {template_info['paragraphs']} paragraf, {template_info['tables']} tabel")
    
    # Analisis styles
    draft_styles = set(draft_info['styles'])
    template_styles = set(template_info['styles'])
    
    missing_styles = template_styles - draft_styles
    extra_styles = draft_styles - template_styles
    
    print(f"\nAnalisis Styles:")
    print(f"  Draft menggunakan    : {len(draft_styles)} styles")
    print(f"  Template menggunakan : {len(template_styles)} styles")
    
    if missing_styles:
        print(f"\n  ❌ Styles yang HILANG di draft (perlu ditambahkan):")
        for style in sorted(missing_styles):
            print(f"     - {style}")
    
    if extra_styles:
        print(f"\n  ⚠️  Styles TAMBAHAN di draft (periksa kesesuaian):")
        for style in sorted(extra_styles):
            print(f"     - {style}")
    
    # Page setup
    if template_specs and 'page_setup' in template_specs:
        print(f"\n{'='*60}")
        print("2. PENGATURAN HALAMAN TEMPLATE")
        print(f"{'='*60}")
        
        ps = template_specs['page_setup']
        print(f"\n  📄 Ukuran Halaman: {ps['page_width']}\" x {ps['page_height']}\" inch")
        print(f"  📏 Margin:")
        print(f"     - Atas   : {ps['top_margin']}\" inch")
        print(f"     - Bawah  : {ps['bottom_margin']}\" inch")
        print(f"     - Kiri   : {ps['left_margin']}\" inch")
        print(f"     - Kanan  : {ps['right_margin']}\" inch")
    
    # Style specifications
    if template_specs and 'styles' in template_specs:
        print(f"\n{'='*60}")
        print("3. SPESIFIKASI STYLES TEMPLATE")
        print(f"{'='*60}")
        
        important_styles = ['Title', 'Heading 1', 'Heading 2', 'Heading 3', 'Normal']
        
        for style_name in important_styles:
            if style_name in template_specs['styles']:
                spec = template_specs['styles'][style_name]
                print(f"\n  📝 {style_name}:")
                
                if 'font_name' in spec:
                    print(f"     Font: {spec['font_name']}")
                if 'font_size' in spec:
                    print(f"     Size: {spec['font_size']} pt")
                if 'bold' in spec:
                    print(f"     Bold: {'Ya' if spec['bold'] else 'Tidak'}")
                if 'italic' in spec:
                    print(f"     Italic: {'Ya' if spec['italic'] else 'Tidak'}")
                if 'alignment' in spec:
                    print(f"     Alignment: {spec['alignment']}")
                if 'space_before' in spec:
                    print(f"     Space Before: {spec['space_before']} pt")
                if 'space_after' in spec:
                    print(f"     Space After: {spec['space_after']} pt")
    
    # Struktur konten
    print(f"\n{'='*60}")
    print("4. STRUKTUR KONTEN")
    print(f"{'='*60}")
    
    print(f"\n  📋 Bagian yang ditemukan di DRAFT:")
    if draft_info['sections']:
        for section_type, text, style in draft_info['sections']:
            print(f"     - {section_type}: '{text[:50]}...' (Style: {style})")
    else:
        print(f"     (Tidak ada bagian khusus yang teridentifikasi)")
    
    print(f"\n  📋 Bagian yang ditemukan di TEMPLATE:")
    if template_info['sections']:
        for section_type, text, style in template_info['sections']:
            print(f"     - {section_type}: '{text[:50]}...' (Style: {style})")
    else:
        print(f"     (Tidak ada bagian khusus yang teridentifikasi)")
    
    # Checklist praktis
    print(f"\n{'='*60}")
    print("5. CHECKLIST SINKRONISASI")
    print(f"{'='*60}")
    
    checklist = [
        "□ Buka draft paper di Microsoft Word",
        "□ Buka template jurnal sebagai referensi",
        "□ Sesuaikan Page Layout (Layout > Margins)",
        f"□ Set margin: Top {ps.get('top_margin', 1)}\" Bottom {ps.get('bottom_margin', 1)}\" Left {ps.get('left_margin', 1)}\" Right {ps.get('right_margin', 1)}\"",
        "□ Copy styles dari template (Home > Styles > Manage Styles)",
        "□ Terapkan style 'Title' untuk judul utama",
        "□ Terapkan style 'Heading 1' untuk bagian utama (Introduction, Method, dll)",
        "□ Terapkan style 'Heading 2' untuk sub-bagian",
        "□ Pastikan semua paragraf menggunakan style 'Normal'",
        "□ Sesuaikan format Abstract dan Keywords",
        "□ Periksa format referensi",
        "□ Review konsistensi font dan spacing",
        "□ Periksa header/footer sesuai template",
        "□ Simpan dokumen dengan format yang benar"
    ]
    
    for item in checklist:
        print(f"   {item}")
    
    # Tips tambahan
    print(f"\n{'='*60}")
    print("6. TIPS PRAKTIS")
    print(f"{'='*60}")
    
    tips = [
        "💡 Gunakan 'Format Painter' untuk menyalin format dengan cepat",
        "💡 Aktifkan 'Show/Hide' (¶) untuk melihat formatting marks",
        "💡 Gunakan 'Find & Replace' untuk mengganti styles secara batch",
        "💡 Simpan template sebagai .dotx untuk penggunaan berulang",
        "💡 Periksa Print Preview sebelum finalisasi",
        "💡 Backup draft asli sebelum melakukan perubahan besar"
    ]
    
    for tip in tips:
        print(f"   {tip}")

# =========================
# FUNGSI OTOMATIS MERGE DRAFT -> TEMPLATE
# =========================

def _iter_block_items(parent):
    """Iterasi item blok (paragraf dan tabel) secara berurutan."""
    parent_elm = parent.element.body if hasattr(parent, 'element') else parent
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _delete_paragraph(paragraph):
    """Hapus paragraf dari dokumen (hack python-docx)."""
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def _clear_document_body(doc: Document):
    """Bersihkan seluruh konten body (paragraf & tabel) dokumen tanpa mengganggu header/footer."""
    # Kumpulkan referensi dulu lalu hapus agar aman
    items = list(_iter_block_items(doc))
    for item in items:
        if isinstance(item, Paragraph):
            try:
                _delete_paragraph(item)
            except Exception:
                pass
        elif isinstance(item, Table):
            try:
                tbl = item._element
                tbl.getparent().remove(tbl)
            except Exception:
                pass
    # Pastikan setidaknya ada satu paragraf kosong
    if len(doc.paragraphs) == 0:
        doc.add_paragraph("")


def _map_style_name(src_style: str, template_doc: Document) -> str:
    """Map nama style dari draft ke style template (heuristik sederhana)."""
    if not src_style:
        return 'Normal'
    name = (src_style or '').strip().lower()
    # Hanya masukkan style yang memiliki nama valid
    template_style_names = {s.name.lower(): s.name for s in template_doc.styles if getattr(s, 'name', None)}

    # Jika style ada persis di template, gunakan langsung
    if name in template_style_names:
        return template_style_names[name]

    # Heuristik umum Heading
    heading_match = re.match(r"heading\s*(\d)", name)
    if heading_match:
        h_level = heading_match.group(1)
        target = f"Heading {h_level}"
        return target if target.lower() in template_style_names else 'Normal'

    # Judul
    if 'title' in name:
        return 'Title' if 'title' in template_style_names else 'Normal'

    # Abstract / Keywords bisa dipetakan ke Normal agar konsisten, nanti user bisa sesuaikan manual
    if any(k in name for k in ['abstract', 'intisari', 'keywords', 'kata kunci']):
        return 'Normal'

    # Default
    return 'Normal'


def _copy_paragraph(src_p: Paragraph, dst_doc: Document, template_doc: Document):
    """Salin paragraf dengan run formatting dasar dan mapping style ke dokumen tujuan."""
    target_style = _map_style_name(getattr(src_p.style, 'name', None), template_doc)
    dst_p = dst_doc.add_paragraph("", style=target_style)

    # Alignment: biarkan style template yang mengatur, tetapi kalau explicit alignment di draft ada, ikutkan
    if src_p.alignment is not None:
        try:
            dst_p.alignment = src_p.alignment
        except Exception:
            pass

    # Salin runs
    for run in src_p.runs:
        new_run = dst_p.add_run(run.text)
        try:
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font and run.font.all_caps is not None:
                new_run.font.all_caps = run.font.all_caps
            if run.font and run.font.small_caps is not None:
                new_run.font.small_caps = run.font.small_caps
            if run.font and run.font.superscript is not None:
                new_run.font.superscript = run.font.superscript
            if run.font and run.font.subscript is not None:
                new_run.font.subscript = run.font.subscript
        except Exception:
            pass


def _copy_table(src_tbl: Table, dst_doc: Document):
    """Salin tabel sederhana (teks sel) ke dokumen tujuan."""
    rows = len(src_tbl.rows)
    cols = len(src_tbl.columns)
    dst_tbl = dst_doc.add_table(rows=rows, cols=cols)
    for r in range(rows):
        for c in range(cols):
            try:
                dst_tbl.cell(r, c).text = src_tbl.cell(r, c).text
            except Exception:
                dst_tbl.cell(r, c).text = ''
    # Biarkan styling tabel mengikuti default template


def merge_draft_into_template(draft_path: str, template_path: str, output_path: str) -> bool:
    """
    Gabungkan isi draft ke dalam dokumen baru berbasis template.
    - Menggunakan style dari template dengan mapping sederhana
    - Menyalin paragraf dan tabel secara berurutan

    Returns: True jika sukses, False jika gagal.
    """
    try:
        draft_doc = Document(draft_path)
        base_template_doc = Document(template_path)

        # Dokumen tujuan berbasis template
        dst_doc = Document(template_path)

        # Bersihkan konten contoh di template agar siap diisi
        _clear_document_body(dst_doc)

        # Iterasi isi draft secara berurutan
        for item in _iter_block_items(draft_doc):
            if isinstance(item, Paragraph):
                _copy_paragraph(item, dst_doc, base_template_doc)
            elif isinstance(item, Table):
                _copy_table(item, dst_doc)

        # Simpan hasil
        dst_doc.save(output_path)
        return True
    except Exception as e:
        print(f"❌ Gagal merge draft ke template: {e}")
        return False

def main():
    """
    Fungsi utama untuk sinkronisasi dokumen
    """
    base_path = "d:/documents/ujaran-kebencian-bahasa-jawa/paper"
    draft_path = os.path.join(base_path, "draft paper.docx")
    template_path = os.path.join(base_path, "New_Template_Jurnal_JITK_Nusa_Mandiri_V11 No 1 August 2025.docx")
    
    print("SINKRONISASI DRAFT PAPER DENGAN TEMPLATE JURNAL")
    print("=" * 80)
    
    # Periksa keberadaan file
    if not os.path.exists(draft_path):
        print(f"❌ Error: Draft paper tidak ditemukan: {draft_path}")
        return
    
    if not os.path.exists(template_path):
        print(f"❌ Error: Template tidak ditemukan: {template_path}")
        return
    
    print(f"✅ Draft paper ditemukan: {os.path.basename(draft_path)}")
    print(f"✅ Template ditemukan: {os.path.basename(template_path)}")
    
    # Analisis dokumen
    print(f"\n🔍 Menganalisis struktur dokumen...")
    
    draft_info = analyze_document_structure(draft_path, "Draft Paper")
    template_info = analyze_document_structure(template_path, "Template Jurnal")
    
    if not draft_info or not template_info:
        print("❌ Error: Gagal menganalisis dokumen")
        return
    
    # Dapatkan spesifikasi template
    print(f"🔍 Menganalisis spesifikasi template...")
    template_specs = get_template_specifications(template_path)
    
    # Generate panduan
    generate_sync_guide(draft_info, template_info, template_specs)

    # Otomatis merge ke dokumen baru berbasis template
    print(f"\n{'='*80}")
    print("7. OTOMATIS MERGE DRAFT -> TEMPLATE")
    print(f"{'='*80}")
    output_path = os.path.join(base_path, "merged_from_draft_into_template.docx")
    print(f"➡️  Membuat dokumen baru: {output_path}")
    ok = merge_draft_into_template(draft_path, template_path, output_path)
    if ok:
        print(f"✅ Berhasil membuat dokumen hasil merge.")
        print(f"📄 Silakan buka dan periksa: {output_path}")
        print("ℹ️ Catatan: Gambar/objek kompleks mungkin perlu penyesuaian manual.")
    else:
        print("❌ Proses merge otomatis gagal. Silakan lakukan penyesuaian manual sesuai panduan di atas.")
    
    print(f"\n{'='*80}")
    print("✅ ANALISIS & MERGE SELESAI")
    print(f"{'='*80}")
    print(f"\n📝 Gunakan panduan di atas untuk menyempurnakan hasil agar sesuai template jurnal.")
    print(f"📁 Backup draft asli sebelum melakukan perubahan lebih jauh.")
    print(f"🔄 Lanjutkan sinkronisasi secara bertahap jika diperlukan.")

if __name__ == "__main__":
    main()